import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import requests
import json
import io
import warnings
warnings.filterwarnings("ignore")


# ── Google Sheets cred system ───────────────────────────────────────────────
from google.oauth2.service_account import Credentials
import gspread

GSHEET_SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]

def _get_sheet():
    """Authenticate and return the credits worksheet."""
    creds_dict = dict(st.secrets["gcp_service_account"])
    creds = Credentials.from_service_account_info(creds_dict, scopes=GSHEET_SCOPES)
    gc = gspread.authorize(creds)
    spreadsheet_id = st.secrets["SHEET_ID"]
    sh = gc.open_by_key(spreadsheet_id)
    return sh.sheet1   # first sheet: Key | Credits | DatePurchased | Email


def lookup_key(access_key: str) -> dict | None:
    """
    Find a row by access key.
    Returns dict with keys: row_index, key, credits, date_purchased, email
    or None if not found.
    """
    try:
        ws = _get_sheet()
        records = ws.get_all_records()           # [{Key, Credits, DatePurchased, Email}, …]
        for i, row in enumerate(records, start=2):  # row 1 is header
            if str(row.get("Key", "")).strip() == access_key.strip():
                return {
                    "row_index": i,
                    "key": row["Key"],
                    "credits": int(row.get("Credits", 0)),
                    "date_purchased": row.get("DatePurchased", ""),
                    "email": row.get("Email", ""),
                }
        return None
    except Exception as e:
        st.error(f"Sheet lookup error: {e}")
        return None


def deduct_credit(row_index: int, current_credits: int) -> int:
    """Write credits − 1 back to the sheet. Returns new credit count."""
    try:
        ws = _get_sheet()
        new_credits = max(0, current_credits - 1)
        # Column B = Credits (column index 2)
        ws.update_cell(row_index, 2, new_credits)
        return new_credits
    except Exception as e:
        st.error(f"Credit deduction error: {e}")
        return current_credits

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="PanelStatX",
    page_icon="⬡",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Custom CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Mono:ital,wght@0,300;0,400;0,500;1,300&family=Syne:wght@400;500;600;700;800&display=swap');

:root {
    --bg:        #0a0c10;
    --surface:   #111318;
    --surface2:  #181c24;
    --border:    #1f2535;
    --accent:    #00e5c8;
    --accent2:   #7c6df0;
    --accent3:   #f05c7c;
    --text:      #e2e8f4;
    --muted:     #6b7a9a;
    --success:   #22d3a0;
    --warn:      #f5a623;
}

html, body, [data-testid="stAppViewContainer"] {
    background-color: var(--bg) !important;
    color: var(--text) !important;
    font-family: 'DM Mono', monospace !important;
}

[data-testid="stSidebar"] {
    background: var(--surface) !important;
    border-right: 1px solid var(--border) !important;
}

/* Headers */
h1, h2, h3, h4 {
    font-family: 'Syne', sans-serif !important;
    color: var(--text) !important;
}

/* Metric cards */
[data-testid="metric-container"] {
    background: var(--surface2) !important;
    border: 1px solid var(--border) !important;
    border-radius: 8px !important;
    padding: 16px !important;
}
[data-testid="metric-container"] > div > div:first-child {
    color: var(--muted) !important;
    font-family: 'DM Mono', monospace !important;
    font-size: 0.72rem !important;
    text-transform: uppercase !important;
    letter-spacing: 0.1em !important;
}
[data-testid="metric-container"] label {
    color: var(--muted) !important;
    font-size: 0.72rem !important;
    text-transform: uppercase !important;
    letter-spacing: 0.1em !important;
}
[data-testid="stMetricValue"] {
    color: var(--accent) !important;
    font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important;
}

/* Buttons */
.stButton > button {
    background: transparent !important;
    border: 1px solid var(--accent) !important;
    color: var(--accent) !important;
    font-family: 'DM Mono', monospace !important;
    font-size: 0.8rem !important;
    border-radius: 4px !important;
    padding: 8px 20px !important;
    transition: all 0.2s !important;
    letter-spacing: 0.05em !important;
}
.stButton > button:hover {
    background: var(--accent) !important;
    color: var(--bg) !important;
}

/* Primary button */
[data-testid="baseButton-primary"] > button,
.stButton [kind="primary"] {
    background: var(--accent) !important;
    color: var(--bg) !important;
    font-weight: 600 !important;
}

/* Selectbox / inputs */
.stSelectbox > div > div,
.stMultiSelect > div > div,
.stTextInput > div > div > input,
.stNumberInput > div > div > input {
    background: var(--surface2) !important;
    border: 1px solid var(--border) !important;
    color: var(--text) !important;
    font-family: 'DM Mono', monospace !important;
    border-radius: 4px !important;
}

/* Tabs */
.stTabs [data-baseweb="tab-list"] {
    background: var(--surface) !important;
    border-bottom: 1px solid var(--border) !important;
    gap: 0 !important;
}
.stTabs [data-baseweb="tab"] {
    color: var(--muted) !important;
    font-family: 'DM Mono', monospace !important;
    font-size: 0.8rem !important;
    background: transparent !important;
    border-radius: 0 !important;
    padding: 12px 24px !important;
}
.stTabs [aria-selected="true"] {
    color: var(--accent) !important;
    border-bottom: 2px solid var(--accent) !important;
    background: transparent !important;
}

/* Dataframes */
.stDataFrame {
    border: 1px solid var(--border) !important;
    border-radius: 8px !important;
}

/* Expanders */
.streamlit-expanderHeader {
    background: var(--surface2) !important;
    color: var(--text) !important;
    font-family: 'DM Mono', monospace !important;
    border: 1px solid var(--border) !important;
    border-radius: 4px !important;
}

/* Sidebar labels */
.stSidebar label, .stSidebar .stMarkdown {
    color: var(--muted) !important;
    font-size: 0.78rem !important;
}

/* Slider */
.stSlider [data-baseweb="slider"] div[role="slider"] {
    background: var(--accent) !important;
}

/* Custom badge */
.badge {
    display: inline-block;
    padding: 2px 10px;
    border-radius: 20px;
    font-size: 0.7rem;
    font-family: 'DM Mono', monospace;
    letter-spacing: 0.05em;
    text-transform: uppercase;
}
.badge-teal  { background: rgba(0,229,200,0.12); color: var(--accent); border: 1px solid rgba(0,229,200,0.3); }
.badge-purple{ background: rgba(124,109,240,0.12); color: var(--accent2); border: 1px solid rgba(124,109,240,0.3); }
.badge-red   { background: rgba(240,92,124,0.12); color: var(--accent3); border: 1px solid rgba(240,92,124,0.3); }
.badge-warn  { background: rgba(245,166,35,0.12); color: var(--warn); border: 1px solid rgba(245,166,35,0.3); }

/* AI box */
.ai-box {
    background: linear-gradient(135deg, rgba(0,229,200,0.05) 0%, rgba(124,109,240,0.08) 100%);
    border: 1px solid rgba(0,229,200,0.25);
    border-left: 3px solid var(--accent);
    border-radius: 8px;
    padding: 20px 24px;
    font-family: 'DM Mono', monospace;
    font-size: 0.85rem;
    line-height: 1.75;
    color: var(--text);
    white-space: pre-wrap;
}
.ai-label {
    font-family: 'Syne', sans-serif;
    font-size: 0.65rem;
    letter-spacing: 0.15em;
    text-transform: uppercase;
    color: var(--accent);
    margin-bottom: 10px;
    display: flex;
    align-items: center;
    gap: 6px;
}
.spinner-dot::after { content: '●'; animation: blink 1s infinite; }
@keyframes blink { 0%,100%{opacity:1} 50%{opacity:0.2} }

/* Hero header */
.hero {
    padding: 28px 0 20px 0;
    border-bottom: 1px solid var(--border);
    margin-bottom: 28px;
}
.hero-title {
    font-family: 'Syne', sans-serif;
    font-size: 2.2rem;
    font-weight: 800;
    letter-spacing: -0.02em;
    color: var(--text);
    margin: 0;
    line-height: 1;
}
.hero-title span { color: var(--accent); }
.hero-sub {
    font-family: 'DM Mono', monospace;
    font-size: 0.78rem;
    color: var(--muted);
    margin-top: 6px;
    letter-spacing: 0.04em;
}
.stat-pill {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    background: var(--surface2);
    border: 1px solid var(--border);
    border-radius: 4px;
    padding: 4px 12px;
    font-size: 0.72rem;
    color: var(--muted);
    margin-right: 8px;
    margin-top: 12px;
    font-family: 'DM Mono', monospace;
}
.stat-pill b { color: var(--accent); }

/* Section card */
.scard {
    background: var(--surface2);
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 20px;
    margin-bottom: 16px;
}
.scard-title {
    font-family: 'Syne', sans-serif;
    font-size: 0.85rem;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    color: var(--muted);
    margin-bottom: 14px;
}

/* Dividers */
hr { border-color: var(--border) !important; }

/* Info boxes */
[data-testid="stInfo"] { background: rgba(0,229,200,0.06) !important; border-left-color: var(--accent) !important; }
[data-testid="stWarning"] { background: rgba(245,166,35,0.06) !important; border-left-color: var(--warn) !important; }
[data-testid="stSuccess"] { background: rgba(34,211,160,0.06) !important; border-left-color: var(--success) !important; }

/* Scrollbar */
::-webkit-scrollbar { width: 4px; height: 4px; }
::-webkit-scrollbar-track { background: var(--bg); }
::-webkit-scrollbar-thumb { background: var(--border); border-radius: 2px; }
</style>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

PLOTLY_THEME = dict(
    paper_bgcolor="rgba(0,0,0,0)",
    plot_bgcolor="rgba(0,0,0,0)",
    font=dict(family="DM Mono, monospace", color="#6b7a9a", size=11),
    xaxis=dict(gridcolor="#1f2535", linecolor="#1f2535", zerolinecolor="#1f2535"),
    yaxis=dict(gridcolor="#1f2535", linecolor="#1f2535", zerolinecolor="#1f2535"),
    colorway=["#00e5c8", "#7c6df0", "#f05c7c", "#f5a623", "#22d3a0", "#60a5fa"],
    margin=dict(l=40, r=20, t=40, b=40),
)


def apply_theme(fig):
    fig.update_layout(**PLOTLY_THEME)
    return fig


def generate_demo_panel():
    """Generate a balanced panel dataset with realistic variation."""
    np.random.seed(42)
    n_entities, n_periods = 30, 10
    entities = [f"Entity_{i:02d}" for i in range(1, n_entities + 1)]
    years = list(range(2014, 2014 + n_periods))
    rows = []
    for e in entities:
        fe = np.random.randn()  # entity fixed effect
        for y in years:
            te = 0.05 * (y - 2014)  # time trend
            x1 = np.random.randn() + fe * 0.3
            x2 = np.random.uniform(0, 10)
            x3 = np.random.choice([0, 1], p=[0.6, 0.4])
            y_val = 2 + 0.8 * x1 - 0.4 * x2 + 1.2 * x3 + fe + te + np.random.randn() * 0.5
            rows.append({"entity": e, "year": y, "y": round(y_val, 4),
                         "x1": round(x1, 4), "x2": round(x2, 4), "x3": int(x3)})
    return pd.DataFrame(rows)


def run_ols(df, y_col, x_cols):
    from numpy.linalg import lstsq
    from scipy import stats as sc_stats
    X = np.column_stack([np.ones(len(df))] + [df[c].values for c in x_cols])
    y = df[y_col].values
    coeffs, residuals, rank, sv = lstsq(X, y, rcond=None)
    y_hat = X @ coeffs
    resid = y - y_hat
    n, k = X.shape
    k_vars = k - 1           
    dof = n - k
    s2 = np.sum(resid**2) / dof
    cov = s2 * np.linalg.inv(X.T @ X)
    se = np.sqrt(np.diag(cov))
    t_stats = coeffs / se
    p_vals = 2 * sc_stats.t.sf(np.abs(t_stats), df=dof)
    ss_tot = np.sum((y - y.mean())**2)
    ss_res = np.sum(resid**2)
    ss_reg = ss_tot - ss_res
    r2 = 1 - ss_res / ss_tot
    r2_adj = 1 - (1 - r2) * (n - 1) / dof
    # F-statistic: (SS_reg / k_vars) / (SS_res / dof)
    f_stat = (ss_reg / max(k_vars, 1)) / (ss_res / max(dof, 1))
    f_p    = 1 - sc_stats.f.cdf(f_stat, dfn=k_vars, dfd=dof)
    names = ["const"] + list(x_cols)
    result_df = pd.DataFrame({"Variable": names, "Coeff": coeffs, "Std_Err": se,
                               "t_stat": t_stats, "p_value": p_vals})
    stats = {"R2": r2, "R2_adj": r2_adj, "N": n, "k": k_vars,
             "AIC": n * np.log(ss_res / n) + 2 * k,
             "BIC": n * np.log(ss_res / n) + k * np.log(n),
             "F_stat": f_stat, "F_p": f_p}
    return result_df, resid, y_hat, stats, cov


def run_within(df, y_col, x_cols, entity_col, time_col):
    """Fixed-effects (within) estimator via demeaning."""
    from scipy import stats as sc_stats
    panel = df.copy()
    for col in [y_col] + list(x_cols):
        entity_means = panel.groupby(entity_col)[col].transform("mean")
        time_means = panel.groupby(time_col)[col].transform("mean")
        grand_mean = panel[col].mean()
        panel[col + "_dm"] = panel[col] - entity_means - time_means + grand_mean
    y_dm = panel[y_col + "_dm"].values
    X_dm = np.column_stack([panel[c + "_dm"].values for c in x_cols])
    from numpy.linalg import lstsq
    coeffs, _, _, _ = lstsq(X_dm, y_dm, rcond=None)
    y_hat_dm = X_dm @ coeffs
    resid = y_dm - y_hat_dm
    n, k = X_dm.shape
    dof = n - k - df[entity_col].nunique() - df[time_col].nunique() + 1
    if dof <= 0:
        dof = max(1, n - k)
    s2 = np.sum(resid**2) / dof
    cov = s2 * np.linalg.inv(X_dm.T @ X_dm)
    se = np.sqrt(np.diag(cov))
    t_stats = coeffs / se
    p_vals = 2 * sc_stats.t.sf(np.abs(t_stats), df=dof)
    ss_tot = np.sum((y_dm - y_dm.mean())**2)
    ss_res = np.sum(resid**2)
    ss_reg = ss_tot - ss_res
    r2 = max(0, 1 - ss_res / ss_tot)
    r2_adj = max(0, 1 - (1 - r2) * (n - 1) / dof)
    f_stat = (ss_reg / max(k, 1)) / (ss_res / max(dof, 1))
    f_p    = 1 - sc_stats.f.cdf(f_stat, dfn=k, dfd=dof)
    result_df = pd.DataFrame({"Variable": list(x_cols), "Coeff": coeffs,
                               "Std_Err": se, "t_stat": t_stats, "p_value": p_vals})
    stats = {"R2": r2, "R2_adj": r2_adj, "N": n, "k": k,
             "AIC": n * np.log(max(ss_res, 1e-10) / n) + 2 * k,
             "BIC": n * np.log(max(ss_res, 1e-10) / n) + k * np.log(n),
             "F_stat": f_stat, "F_p": f_p}
    return result_df, resid, y_hat_dm, stats, cov


def run_re(df, y_col, x_cols, entity_col, time_col):
    """
    Random Effects estimator (Swamy-Arora / GLS).
    Estimates the between-entity variance (sigma_u²) and within-entity
    variance (sigma_e²), computes the GLS theta weight, then runs quasi-
    demeaned OLS to obtain RE coefficients, SEs, t-stats, and p-values.
    Also returns the coefficient vcov for use in the Hausman test.
    """
    from numpy.linalg import lstsq
    from scipy import stats as sc_stats

    panel = df.copy().sort_values([entity_col, time_col])
    n_entities = panel[entity_col].nunique()
    T = panel[time_col].nunique()          # assume balanced panel
    N = len(panel)
    k = len(x_cols)

    # ── Step 1: within (FE) residuals to estimate sigma_e² ────────────────────
    result_fe, resid_fe, _, stats_fe, _ = run_within(
        panel, y_col, x_cols, entity_col, time_col
    )
    dof_fe = max(N - k - n_entities, 1)
    sigma_e2 = np.sum(resid_fe ** 2) / dof_fe

    # ── Step 2: between residuals to estimate sigma_u² ────────────────────────
    grp = panel.groupby(entity_col)[[y_col] + list(x_cols)].mean().reset_index()
    y_b  = grp[y_col].values
    X_b  = np.column_stack([np.ones(n_entities)] + [grp[c].values for c in x_cols])
    b_coeffs, _, _, _ = lstsq(X_b, y_b, rcond=None)
    resid_b = y_b - X_b @ b_coeffs
    sigma_b2 = max(0.0, np.sum(resid_b ** 2) / max(n_entities - k - 1, 1) - sigma_e2 / T)
    sigma_u2 = sigma_b2

    # ── Step 3: GLS theta weight ───────────────────────────────────────────────
    theta = 1.0 - np.sqrt(sigma_e2 / max(T * sigma_u2 + sigma_e2, 1e-12))

    # ── Step 4: quasi-demean (partial within) ─────────────────────────────────
    panel2 = panel.copy()
    for col in [y_col] + list(x_cols):
        entity_mean = panel2.groupby(entity_col)[col].transform("mean")
        panel2[col + "_qd"] = panel2[col] - theta * entity_mean

    y_qd = panel2[y_col + "_qd"].values
    X_qd = np.column_stack([np.ones(N)] + [panel2[c + "_qd"].values for c in x_cols])

    # ── Step 5: OLS on quasi-demeaned data ────────────────────────────────────
    coeffs, _, _, _ = lstsq(X_qd, y_qd, rcond=None)
    y_hat  = X_qd @ coeffs
    resid  = y_qd  - y_hat
    dof    = max(N - k - 1, 1)
    s2     = np.sum(resid ** 2) / dof
    cov    = s2 * np.linalg.inv(X_qd.T @ X_qd)
    se     = np.sqrt(np.diag(cov))
    t_stats = coeffs / se
    p_vals  = 2 * sc_stats.t.sf(np.abs(t_stats), df=dof)

    ss_tot = np.sum((y_qd - y_qd.mean()) ** 2)
    ss_res = np.sum(resid ** 2)
    ss_reg = ss_tot - ss_res
    r2     = max(0.0, 1 - ss_res / max(ss_tot, 1e-12))
    r2_adj = max(0.0, 1 - (1 - r2) * (N - 1) / dof)
    f_stat = (ss_reg / max(k, 1)) / (ss_res / dof)
    f_p    = 1 - sc_stats.f.cdf(f_stat, dfn=k, dfd=dof)

    names = ["const"] + list(x_cols)
    result_df = pd.DataFrame({"Variable": names, "Coeff": coeffs, "Std_Err": se,
                               "t_stat": t_stats, "p_value": p_vals})
    stats = {
        "R2": r2, "R2_adj": r2_adj, "N": N, "k": k,
        "AIC": N * np.log(max(ss_res, 1e-10) / N) + 2 * (k + 1),
        "BIC": N * np.log(max(ss_res, 1e-10) / N) + (k + 1) * np.log(N),
        "F_stat": f_stat, "F_p": f_p,
        "sigma_u2": sigma_u2, "sigma_e2": sigma_e2, "theta": theta,
    }
    return result_df, resid, y_hat, stats, cov


def run_fd(df, y_col, x_cols, entity_col, time_col):
    """First-difference estimator."""
    panel = df.sort_values([entity_col, time_col]).copy()
    fd = panel.groupby(entity_col)[[y_col] + list(x_cols)].diff().dropna()
    return run_ols(fd, y_col, x_cols)


def breusch_pagan_test(resid, X):
    """
    Breusch-Pagan / Cook-Weisberg test for heteroskedasticity.
    Regresses squared residuals on the regressors X; the LM statistic
    is n * R² of that auxiliary regression, chi-sq(k) distributed.
    X should include the intercept column.
    """
    from scipy import stats as sc_stats
    from numpy.linalg import lstsq
    resid = np.asarray(resid, dtype=float)
    e2 = resid ** 2
    # auxiliary regression of e² on X
    coeffs_aux, _, _, _ = lstsq(X, e2, rcond=None)
    e2_hat = X @ coeffs_aux
    ss_tot_aux = np.sum((e2 - e2.mean()) ** 2)
    ss_res_aux = np.sum((e2 - e2_hat) ** 2)
    r2_aux = max(0.0, 1 - ss_res_aux / max(ss_tot_aux, 1e-12))
    n = len(resid)
    k = X.shape[1] - 1          # exclude intercept
    bp_stat = n * r2_aux
    bp_p    = 1 - sc_stats.chi2.cdf(bp_stat, df=k)
    return bp_stat, bp_p, k


    """Simple Hausman test statistic."""
    diff = fe_coef - re_coef
    diff_vcov = fe_vcov - re_vcov
    try:
        stat = float(diff @ np.linalg.inv(diff_vcov) @ diff)
        df = len(diff)
        from scipy import stats as sc_stats
        p = 1 - sc_stats.chi2.cdf(stat, df)
        return stat, p, df
    except Exception:
        return None, None, None


def significance_stars(p):
    if p < 0.001: return "***"
    if p < 0.01:  return "**"
    if p < 0.05:  return "*"
    if p < 0.1:   return "·"
    return ""


def call_openai(system_prompt, user_prompt):
    """Call AI Model API for AI explanations."""
    try:
        api_key = st.secrets["OPENAI_API_KEY"]
    except Exception:
        return "OpenAI API key not configured or Exhausted."

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    payload = {
        "model": "gpt-4o",
        "max_tokens": 1000,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
    }
    try:
        resp = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers, json=payload, timeout=30,
        )
        data = resp.json()
        if "choices" in data:
            return data["choices"][0]["message"]["content"]
        return f"API error: {data.get('error', {}).get('message', str(data))}"
    except Exception as e:
        return f"Request failed: {e}"


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX REPORT GENERATOR
# ═══════════════════════════════════════════════════════════════════════════════

def build_docx_report(res, model_type, ai_explanation=""):
    """
    Build a professional Word document containing:
      • Cover page
      • Model summary & fit statistics table
      • Coefficient estimates table
      • Residual diagnostics table
      • AI write-up (if available)
    Returns bytes suitable for st.download_button.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime, io, numpy as np
    from scipy import stats as sc_stats

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # ── Helper: set cell shading ───────────────────────────────────────────────
    def shade_cell(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    def set_cell_border(cell, **kwargs):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            if edge in kwargs:
                tag  = OxmlElement(f'w:{edge}')
                tag.set(qn('w:val'),   kwargs[edge].get('val', 'single'))
                tag.set(qn('w:sz'),    str(kwargs[edge].get('sz', 4)))
                tag.set(qn('w:color'), kwargs[edge].get('color', '000000'))
                tcBorders.append(tag)
        tcPr.append(tcBorders)

    # ── Colour palette ────────────────────────────────────────────────────────
    DARK_BG   = "0A0C10"
    ACCENT    = "00C8B0"   # teal (print-safe, slightly darker)
    HEADER_BG = "1A2035"
    ALT_ROW   = "F4F7FA"
    WHITE     = "FFFFFF"
    TEXT_DARK = RGBColor(0x1A, 0x20, 0x35)
    TEAL_RGB  = RGBColor(0x00, 0xC8, 0xB0)
    GRAY_RGB  = RGBColor(0x6B, 0x7A, 0x9A)

    # ── Font helpers ──────────────────────────────────────────────────────────
    def h_style(para, text, size=14, bold=True, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6):
        para.alignment = align
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after  = Pt(space_after)
        run = para.add_run(text)
        run.bold       = bold
        run.font.size  = Pt(size)
        run.font.name  = "Arial"
        if color:
            run.font.color.rgb = color
        return run

    def body_para(doc, text, size=10, color=None, bold=False, italic=False, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.name = "Arial"
        r.bold      = bold
        r.italic    = italic
        if color:
            r.font.color.rgb = color
        return p

    def add_rule(doc, color_hex="D0D8E8", thickness=6):
        """Thin horizontal rule via paragraph border."""
        p    = doc.add_paragraph()
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    str(thickness))
        bot.set(qn('w:color'), color_hex)
        bot.set(qn('w:space'), '1')
        pBdr.append(bot)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    # Dark-header block via a 1-row table
    cover_tbl = doc.add_table(rows=1, cols=1)
    cover_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cover_cell = cover_tbl.rows[0].cells[0]
    shade_cell(cover_cell, DARK_BG)
    cover_cell.width = Inches(6.5)
    cover_cell._tc.get_or_add_tcPr()

    cp = cover_cell.add_paragraph()
    cp.paragraph_format.space_before = Pt(18)
    cp.paragraph_format.space_after  = Pt(2)
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = cp.add_run("⬡  PanelStatX")
    r1.font.size = Pt(22)
    r1.font.bold = True
    r1.font.name = "Arial"
    r1.font.color.rgb = TEAL_RGB

    cp2 = cover_cell.add_paragraph()
    cp2.paragraph_format.space_before = Pt(2)
    cp2.paragraph_format.space_after  = Pt(4)
    cp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = cp2.add_run("Panel Regression Analysis Report")
    r2.font.size  = Pt(14)
    r2.font.name  = "Arial"
    r2.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF4)

    cp3 = cover_cell.add_paragraph()
    cp3.paragraph_format.space_before = Pt(0)
    cp3.paragraph_format.space_after  = Pt(18)
    cp3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    now_str = datetime.datetime.now().strftime("%d %B %Y, %H:%M")
    r3 = cp3.add_run(f"Generated: {now_str}   ·   Estimator: {model_type}")
    r3.font.size  = Pt(9)
    r3.font.name  = "Arial"
    r3.font.color.rgb = RGBColor(0x6B, 0x7A, 0x9A)

    doc.add_paragraph()  # spacer

    # Cover meta row
    result_df = res["result_df"]
    stats     = res["stats"]
    resid     = np.asarray(res["resid"], dtype=float)
    resid     = resid[np.isfinite(resid)]
    y_col     = res["y_col"]
    x_cols    = res["x_cols"]
    entity_col = res["entity_col"]
    time_col   = res["time_col"]

    meta_items = [
        ("Dependent Variable", y_col),
        ("Independent Variables", ", ".join(x_cols)),
        ("Entity Column", entity_col),
        ("Time Column", time_col),
        ("Observations (N)", f"{stats['N']:,}"),
        ("Variables (k)", str(stats["k"])),
    ]
    meta_tbl = doc.add_table(rows=len(meta_items), cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_widths = [Inches(2.2), Inches(4.3)]
    for i, (label, value) in enumerate(meta_items):
        row = meta_tbl.rows[i]
        lc, vc = row.cells[0], row.cells[1]
        lc.width = col_widths[0]
        vc.width = col_widths[1]
        for edge in ('top', 'bottom', 'left', 'right'):
            set_cell_border(lc, **{edge: {'val': 'none', 'sz': 0, 'color': 'FFFFFF'}})
            set_cell_border(vc, **{edge: {'val': 'none', 'sz': 0, 'color': 'FFFFFF'}})
        lp = lc.add_paragraph(label)
        lp.runs[0].font.size = Pt(9)
        lp.runs[0].font.name = "Arial"
        lp.runs[0].bold = True
        lp.runs[0].font.color.rgb = GRAY_RGB
        vp = vc.add_paragraph(value)
        vp.runs[0].font.size = Pt(9)
        vp.runs[0].font.name = "Arial"
        vp.runs[0].font.color.rgb = TEXT_DARK

    doc.add_page_break()

    # ── Helper: set thick/thin academic table borders ─────────────────────────
    def set_tbl_cell_border_academic(cell, top=None, bottom=None, left=None, right=None):
        """Set individual cell borders for SPSS/Stata-style academic tables."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge, spec in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            tag = OxmlElement(f'w:{edge}')
            if spec:
                tag.set(qn('w:val'),   spec.get('val', 'single'))
                tag.set(qn('w:sz'),    str(spec.get('sz', 4)))
                tag.set(qn('w:color'), spec.get('color', '000000'))
            else:
                tag.set(qn('w:val'), 'none')
                tag.set(qn('w:sz'),  '0')
                tag.set(qn('w:color'), 'FFFFFF')
            tcBorders.append(tag)
        tcPr.append(tcBorders)

    THICK_BORDER  = {'val': 'single', 'sz': 12, 'color': '1A2035'}  # 1.5pt top/bottom rule
    THIN_BORDER   = {'val': 'single', 'sz':  4, 'color': '9AA3BE'}  # 0.5pt header underline
    NO_BORDER     = None

    def _cell(row_obj, j, text, bold=False, italic=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT,
              color=None, top=None, bottom=None, width=None):
        c = row_obj.cells[j]
        if width:
            c.width = width
        p = c.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(text)
        r.font.size  = Pt(size)
        r.font.name  = "Arial"
        r.bold       = bold
        r.italic     = italic
        if color:
            r.font.color.rgb = color
        set_tbl_cell_border_academic(c, top=top, bottom=bottom, left=NO_BORDER, right=NO_BORDER)
        return c

    # ── SECTION 1: MODEL FIT STATISTICS  (SPSS/Stata style) ──────────────────
    s1 = doc.add_paragraph()
    h_style(s1, "Table 1.  Model Fit Statistics", size=12, color=TEXT_DARK, space_before=0, bold=True)

    # Caption line
    cap1 = doc.add_paragraph()
    cap1.paragraph_format.space_before = Pt(0)
    cap1.paragraph_format.space_after  = Pt(4)
    cap1_r = cap1.add_run(f"Dependent Variable: {res['y_col']}   |   Estimator: {model_type}   |   N = {stats['N']:,}")
    cap1_r.font.size  = Pt(8.5)
    cap1_r.font.name  = "Arial"
    cap1_r.italic     = True
    cap1_r.font.color.rgb = GRAY_RGB

    # Two-column layout: left = stat name, right = value (right-aligned numbers)
    fit_rows = [
        ("R²",                         f"{stats['R2']:.4f}"),
        ("Adjusted R²",                f"{stats['R2_adj']:.4f}"),
        ("F-statistic",                f"{stats.get('F_stat', float('nan')):.4f}"),
        ("Prob (F-statistic)",         f"{stats.get('F_p', float('nan')):.4f}"),
        ("Akaike Info. Criterion (AIC)", f"{stats['AIC']:.2f}"),
        ("Bayesian Info. Criterion (BIC)", f"{stats['BIC']:.2f}"),
        ("No. of Observations",        f"{stats['N']:,}"),
        ("No. of Regressors (k)",      f"{stats['k']}"),
    ]

    # Optionally add RE-specific rows
    if stats.get("sigma_u2") is not None:
        fit_rows += [
            ("Variance: Between-Entity (σ²ᵤ)", f"{stats['sigma_u2']:.6f}"),
            ("Variance: Within-Entity (σ²ₑ)",  f"{stats['sigma_e2']:.6f}"),
            ("GLS Theta (θ)",                    f"{stats['theta']:.4f}"),
        ]

    n_fit = len(fit_rows)
    COL_W_STAT = Inches(3.8)
    COL_W_VAL  = Inches(2.7)


    def clear_tbl_borders(tbl):
        """Remove all table-level borders so only cell-level borders show."""
        tbl_el  = tbl._tbl
        tblPr   = tbl_el.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl_el.insert(0, tblPr)
        for existing in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(existing)
        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"),   "none")
            tag.set(qn("w:sz"),    "0")
            tag.set(qn("w:color"), "FFFFFF")
            tblBorders.append(tag)
        tblPr.append(tblBorders)

    fit_tbl = doc.add_table(rows=n_fit + 2, cols=2)  # +2: header row + bottom rule row
    fit_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    fit_tbl.style     = "Table Grid"
    clear_tbl_borders(fit_tbl)

    # --- Top rule row (row 0) ---
    for j in range(2):
        c = fit_tbl.rows[0].cells[j]
        c.width = COL_W_STAT if j == 0 else COL_W_VAL
        p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        set_tbl_cell_border_academic(c, top=THICK_BORDER, bottom=THIN_BORDER, left=NO_BORDER, right=NO_BORDER)
        r = p.add_run("Statistic" if j == 0 else "Value")
        r.font.size = Pt(9); r.font.name = "Arial"; r.bold = True
        r.font.color.rgb = TEXT_DARK
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if j == 1 else WD_ALIGN_PARAGRAPH.LEFT

    # --- Data rows ---
    for i, (label, value) in enumerate(fit_rows):
        row_obj = fit_tbl.rows[i + 1]
        is_last = (i == n_fit - 1)
        bot = THIN_BORDER if is_last else NO_BORDER
        _cell(row_obj, 0, label, bold=False, size=9, color=TEXT_DARK,
              top=NO_BORDER, bottom=bot, width=COL_W_STAT)
        _cell(row_obj, 1, value, bold=True, size=9, color=TEXT_DARK,
              align=WD_ALIGN_PARAGRAPH.RIGHT, top=NO_BORDER, bottom=bot, width=COL_W_VAL)

    # --- Bottom rule row ---
    for j in range(2):
        c = fit_tbl.rows[n_fit + 1].cells[j]
        c.width = COL_W_STAT if j == 0 else COL_W_VAL
        p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        set_tbl_cell_border_academic(c, top=THICK_BORDER, bottom=NO_BORDER, left=NO_BORDER, right=NO_BORDER)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── SECTION 2: COEFFICIENT ESTIMATES TABLE  (Stata style) ───────────────
    s2 = doc.add_paragraph()
    h_style(s2, "Table 2.  Coefficient Estimates", size=12, color=TEXT_DARK, space_before=12, bold=True)

    cap2 = doc.add_paragraph()
    cap2.paragraph_format.space_before = Pt(0)
    cap2.paragraph_format.space_after  = Pt(4)
    cap2_r = cap2.add_run(
        f"Dependent Variable: {res['y_col']}   |   Estimator: {model_type}   |   "
        f"R² = {stats['R2']:.4f}   |   F({stats['k']}, {stats['N'] - stats['k'] - 1}) = "
        f"{stats.get('F_stat', float('nan')):.2f}   |   Prob > F = {stats.get('F_p', float('nan')):.4f}"
    )
    cap2_r.font.size = Pt(8.5); cap2_r.font.name = "Arial"
    cap2_r.italic = True; cap2_r.font.color.rgb = GRAY_RGB

    # Headers: Variable | Coef. | Std. Err. | t | P>|t| | [95% Conf. Interval]
    # Merge last two header cells conceptually — do it with two separate cells
    coef_headers = ["Variable", "Coef.", "Std. Err.", "t", "P>|t|", "[95% Conf.", "Interval]"]
    COL_WC = [Inches(1.5), Inches(0.85), Inches(0.85), Inches(0.7), Inches(0.7), Inches(0.85), Inches(0.85)]

    n_coef_rows = len(result_df)
    coef_tbl = doc.add_table(rows=n_coef_rows + 2, cols=len(coef_headers))
    coef_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    coef_tbl.style     = "Table Grid"
    clear_tbl_borders(coef_tbl)

    # --- Header row (row 0): thick top + thin bottom ---
    for j, (h, w) in enumerate(zip(coef_headers, COL_WC)):
        c = coef_tbl.rows[0].cells[j]
        c.width = w
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        aln = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        p.alignment = aln
        r = p.add_run(h); r.font.bold = True; r.font.size = Pt(9); r.font.name = "Arial"
        r.font.color.rgb = TEXT_DARK
        set_tbl_cell_border_academic(c, top=THICK_BORDER, bottom=THIN_BORDER, left=NO_BORDER, right=NO_BORDER)

    # --- Data rows ---
    for i, row_data in result_df.iterrows():
        p_val  = row_data["p_value"]
        stars  = significance_stars(p_val)
        ci_lo  = row_data["Coeff"] - 1.96 * row_data["Std_Err"]
        ci_hi  = row_data["Coeff"] + 1.96 * row_data["Std_Err"]
        is_last = (i == n_coef_rows - 1)
        bot = THIN_BORDER if is_last else NO_BORDER

        # Coefficient cell: value + stars appended (like Stata)
        coef_str  = f"{row_data['Coeff']:.4f}{stars}"
        row_vals  = [
            row_data["Variable"],
            coef_str,
            f"{row_data['Std_Err']:.4f}",
            f"{row_data['t_stat']:.3f}",
            f"{p_val:.4f}",
            f"{ci_lo:.4f}",
            f"{ci_hi:.4f}",
        ]
        tbl_row = coef_tbl.rows[i + 1]
        for j, (val, w) in enumerate(zip(row_vals, COL_WC)):
            c = tbl_row.cells[j]
            c.width = w
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(9); r.font.name = "Arial"
            r.bold = (j == 0)
            r.font.color.rgb = TEXT_DARK
            set_tbl_cell_border_academic(c, top=NO_BORDER, bottom=bot, left=NO_BORDER, right=NO_BORDER)

    # --- Bottom thick rule row ---
    for j in range(len(coef_headers)):
        c = coef_tbl.rows[n_coef_rows + 1].cells[j]
        c.width = COL_WC[j]
        p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        set_tbl_cell_border_academic(c, top=THICK_BORDER, bottom=NO_BORDER, left=NO_BORDER, right=NO_BORDER)

    # --- Table footnote ---
    fn_p = doc.add_paragraph()
    fn_p.paragraph_format.space_before = Pt(3)
    fn_p.paragraph_format.space_after  = Pt(8)
    fn_r = fn_p.add_run(
        "Notes: Standard errors in parentheses.  "
        "*** p<0.001  ** p<0.01  * p<0.05  \u00b7 p<0.1  "
        "Confidence intervals computed at the 95% level (±1.96 × SE)."
    )
    fn_r.font.size = Pt(7.5); fn_r.font.name = "Arial"; fn_r.italic = True
    fn_r.font.color.rgb = GRAY_RGB

    # ── SECTION 3: RESIDUAL DIAGNOSTICS  (academic style) ────────────────────
    doc.add_page_break()
    s3 = doc.add_paragraph()
    h_style(s3, "Table 3.  Residual Diagnostics & Specification Tests", size=12, color=TEXT_DARK, space_before=0, bold=True)

    cap3 = doc.add_paragraph()
    cap3.paragraph_format.space_before = Pt(0)
    cap3.paragraph_format.space_after  = Pt(4)
    cap3_r = cap3.add_run("Post-estimation residual analysis and hypothesis tests for model validity.")
    cap3_r.font.size = Pt(8.5); cap3_r.font.name = "Arial"
    cap3_r.italic = True; cap3_r.font.color.rgb = GRAY_RGB

    jb_stat, jb_p = sc_stats.jarque_bera(resid)
    dw_stat        = np.sum(np.diff(resid)**2) / max(np.sum(resid**2), 1e-10)
    skewness       = sc_stats.skew(resid)
    kurt           = sc_stats.kurtosis(resid)

    bp_stat = stats.get("BP_stat")
    bp_p    = stats.get("BP_p")
    bp_stat_str = f"{bp_stat:.4f}" if bp_stat is not None else "N/A"
    bp_p_str    = f"{bp_p:.4f}"    if bp_p    is not None else "N/A"

    # Interpretations
    normality_verdict = "H\u2080 not rejected (normality)" if jb_p >= 0.05 else "H\u2080 rejected \u2014 non-normal residuals"
    dw_verdict        = "No autocorrelation detected" if 1.5 <= dw_stat <= 2.5 else "Possible autocorrelation"
    bp_verdict        = ("H\u2080 not rejected (homoskedastic)" if bp_p is not None and bp_p >= 0.05
                         else ("H\u2080 rejected \u2014 heteroskedasticity detected" if bp_p is not None else "Not computed"))
    skew_verdict      = "Symmetric" if abs(skewness) < 0.5 else ("Moderate skew" if abs(skewness) < 1 else "High skew")
    kurt_verdict      = "Mesokurtic" if abs(kurt) < 1 else ("Leptokurtic (fat tails)" if kurt > 1 else "Platykurtic (thin tails)")

    # Two-panel layout: (A) Residual Summary  (B) Test Statistics
    diag_left = [
        ("Panel A: Residual Summary", "", ""),
        ("Mean of Residuals",       f"{np.mean(resid):.6f}",   "Near zero \u21d2 unbiased"),
        ("Std. Dev. of Residuals",  f"{np.std(resid):.6f}",    "Overall residual spread"),
        ("Min. Residual",           f"{np.min(resid):.6f}",    "Largest negative deviation"),
        ("Max. Residual",           f"{np.max(resid):.6f}",    "Largest positive deviation"),
        ("Skewness",                f"{skewness:.4f}",          skew_verdict),
        ("Excess Kurtosis",         f"{kurt:.4f}",              kurt_verdict),
        ("", "", ""),
        ("Panel B: Specification Tests", "", ""),
        ("Jarque-Bera Statistic",   f"{jb_stat:.4f}",           "H\u2080: Residuals are normally distributed"),
        ("Jarque-Bera p-value",     f"{jb_p:.4f}",              normality_verdict),
        ("Durbin-Watson Statistic", f"{dw_stat:.4f}",           dw_verdict + " (Benchmark: 1.5\u20132.5)"),
        ("Breusch-Pagan LM Stat.",  bp_stat_str,                "H\u2080: Constant error variance"),
        ("Breusch-Pagan p-value",   bp_p_str,                   bp_verdict),
    ]

    COL_D = [Inches(2.5), Inches(1.2), Inches(2.8)]
    n_diag = len(diag_left)
    diag_tbl = doc.add_table(rows=n_diag + 2, cols=3)
    diag_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    diag_tbl.style     = "Table Grid"
    clear_tbl_borders(diag_tbl)

    # Header row
    for j, (hdr, w) in enumerate(zip(["Test / Statistic", "Value", "Decision / Notes"], COL_D)):
        c = diag_tbl.rows[0].cells[j]
        c.width = w
        p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j == 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(hdr); r.font.bold = True; r.font.size = Pt(9); r.font.name = "Arial"
        r.font.color.rgb = TEXT_DARK
        set_tbl_cell_border_academic(c, top=THICK_BORDER, bottom=THIN_BORDER, left=NO_BORDER, right=NO_BORDER)

    for i, (d_name, d_val, d_note) in enumerate(diag_left):
        row_obj = diag_tbl.rows[i + 1]
        is_panel_hdr = d_name.startswith("Panel")
        is_last      = (i == n_diag - 1)
        is_blank     = (d_name == "")
        bot          = THIN_BORDER if is_last else NO_BORDER

        for j, (txt, w) in enumerate(zip([d_name, d_val, d_note], COL_D)):
            c = row_obj.cells[j]
            c.width = w
            p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j == 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(txt)
            r.font.size = Pt(9 if not is_panel_hdr else 8.5)
            r.font.name = "Arial"
            r.bold      = is_panel_hdr or (j == 0 and not is_blank)
            r.italic    = is_panel_hdr
            r.font.color.rgb = GRAY_RGB if is_panel_hdr else TEXT_DARK
            set_tbl_cell_border_academic(c, top=NO_BORDER, bottom=bot, left=NO_BORDER, right=NO_BORDER)

    # Bottom thick rule
    for j in range(3):
        c = diag_tbl.rows[n_diag + 1].cells[j]
        c.width = COL_D[j]
        p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        set_tbl_cell_border_academic(c, top=THICK_BORDER, bottom=NO_BORDER, left=NO_BORDER, right=NO_BORDER)

    fn3 = doc.add_paragraph()
    fn3.paragraph_format.space_before = Pt(3); fn3.paragraph_format.space_after = Pt(8)
    fn3_r = fn3.add_run(
        "Notes: Jarque-Bera tests H\u2080: residuals follow a normal distribution (chi-sq, df=2).  "
        "Breusch-Pagan tests H\u2080: homoskedastic errors via auxiliary LM regression (chi-sq, df=k).  "
        "Durbin-Watson near 2 indicates absence of first-order autocorrelation."
    )
    fn3_r.font.size = Pt(7.5); fn3_r.font.name = "Arial"; fn3_r.italic = True
    fn3_r.font.color.rgb = GRAY_RGB

    # ── SECTION 4: AI WRITE-UP  (rich formatted) ──────────────────────────────
    if ai_explanation and ai_explanation.strip():
        doc.add_page_break()
        s4 = doc.add_paragraph()
        h_style(s4, "4.  AI Econometric Interpretation", size=12, color=TEXT_DARK, space_before=0, bold=True)
        add_rule(doc)

        body_para(doc,
            "The following structured interpretation was generated by an AI model based on the regression output above. "
            "Statistical values are taken directly from the tables above; commentary reflects econometric best-practice.",
            size=8, color=GRAY_RGB, italic=True, space_after=10)

        import re as _re

        def _render_inline(para_obj, text):
            """Render **bold** and *italic* inline markers into a Word paragraph."""
            # Split on **...** and *...* markers
            parts = _re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = para_obj.add_run(part[2:-2])
                    r.bold = True; r.font.size = Pt(9.5); r.font.name = "Arial"; r.font.color.rgb = TEXT_DARK
                elif part.startswith('*') and part.endswith('*'):
                    r = para_obj.add_run(part[1:-1])
                    r.italic = True; r.font.size = Pt(9.5); r.font.name = "Arial"; r.font.color.rgb = TEXT_DARK
                else:
                    r = para_obj.add_run(part)
                    r.font.size = Pt(9.5); r.font.name = "Arial"; r.font.color.rgb = TEXT_DARK

        for line in ai_explanation.strip().split("\n"):
            stripped = line.strip()
            if not stripped:
                sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(3)
                continue

            # Section heading: ## or ### or lines ending with a colon that are short (≤60 chars)
            if stripped.startswith("## ") or stripped.startswith("### "):
                hdr_txt = stripped.lstrip("# ").rstrip(":")
                hp = doc.add_paragraph()
                hp.paragraph_format.space_before = Pt(10)
                hp.paragraph_format.space_after  = Pt(3)
                hr = hp.add_run(hdr_txt.upper())
                hr.font.size = Pt(9); hr.font.name = "Arial"; hr.bold = True
                hr.font.color.rgb = GRAY_RGB
                # Add thin underline via paragraph border
                pPr  = hp._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bot  = OxmlElement('w:bottom')
                bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
                bot.set(qn('w:color'), 'D0D8E8'); bot.set(qn('w:space'), '1')
                pBdr.append(bot); pPr.append(pBdr)
                continue

            # Numbered list:  "1. " or "1) "
            m_num = _re.match(r'^(\d+)[.)]\s+(.*)', stripped)
            if m_num:
                p = doc.add_paragraph()
                p.paragraph_format.space_after  = Pt(4)
                p.paragraph_format.left_indent  = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                nr = p.add_run(f"{m_num.group(1)}.  ")
                nr.font.size = Pt(9.5); nr.font.name = "Arial"; nr.bold = True; nr.font.color.rgb = GRAY_RGB
                _render_inline(p, m_num.group(2))
                continue

            # Bullet list: "- " or "• "
            m_bul = _re.match(r'^[-\u2022]\s+(.*)', stripped)
            if m_bul:
                p = doc.add_paragraph()
                p.paragraph_format.space_after  = Pt(3)
                p.paragraph_format.left_indent  = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                br = p.add_run("\u2013  ")
                br.font.size = Pt(9.5); br.font.name = "Arial"; br.bold = True; br.font.color.rgb = GRAY_RGB
                _render_inline(p, m_bul.group(1))
                continue

            # Bold label lines: "Label: text"
            m_lbl = _re.match(r'^(\*\*[^*]+\*\*|[A-Z][^:]{2,40}):\s+(.*)', stripped)
            if m_lbl and not stripped.startswith("http"):
                p = doc.add_paragraph()
                p.paragraph_format.space_after  = Pt(4)
                p.paragraph_format.left_indent  = Inches(0.1)
                lr = p.add_run(m_lbl.group(1).strip('*') + ": ")
                lr.font.size = Pt(9.5); lr.font.name = "Arial"; lr.bold = True; lr.font.color.rgb = TEXT_DARK
                _render_inline(p, m_lbl.group(2))
                continue

            # Normal paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(5)
            p.paragraph_format.left_indent  = Inches(0.1)
            _render_inline(p, stripped)

    else:
        s4 = doc.add_paragraph()
        h_style(s4, "4.  AI Econometric Interpretation", size=12, color=TEXT_DARK, space_before=6, bold=True)
        add_rule(doc)
        body_para(doc,
            "No AI interpretation has been generated for this session. "
            "Run the AI Explainer from the AI Explainer tab, then download the report again to include it.",
            size=9, color=GRAY_RGB, italic=True)

    # ── FOOTER on every page ──────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(f"PanelStatX  ·  Panel Regression Analysis Report  ·  {datetime.datetime.now().strftime('%Y-%m-%d')}")
        fr.font.size = Pt(7)
        fr.font.name = "Arial"
        fr.font.color.rgb = GRAY_RGB

    # ── Serialise to bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ═══════════════════════════════════════════════════════════════════════════════

for key, default in [
    ("df", None), ("results", None), ("ai_explanation", ""),
    ("model_type", "Fixed Effects (Two-Way)"),
    ("access_granted", False), ("access_error", ""),
    # Credit-system state
    ("user_key", ""), ("user_credits", 0),
    ("user_email", ""), ("user_row", None),
]:
    if key not in st.session_state:
        st.session_state[key] = default

# ═══════════════════════════════════════════════════════════════════════════════
# ACCESS KEY GATE 
# ═══════════════════════════════════════════════════════════════════════════════

if not st.session_state.access_granted:

    # ── LANDING PAGE CSS ──────────────────────────────────────────────────────
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=DM+Mono:ital,wght@0,300;0,400;0,500&family=Bricolage+Grotesque:opsz,wght@12..96,300;12..96,400;12..96,500;12..96,600;12..96,700;12..96,800&display=swap');

    /* ── Global reset for landing ── */
    [data-testid="stSidebar"]                 { display: none !important; }
    [data-testid="stSidebarCollapsedControl"]  { display: none !important; }
    [data-testid="stAppViewContainer"]         { padding: 0 !important; background: #05070f !important; }
    [data-testid="block-container"]            { padding: 0 !important; max-width: 100% !important; }
    section.main > div                         { padding: 0 !important; }
    html, body                                 { background: #05070f !important; }
    footer, #MainMenu, [data-testid="stToolbar"] { display: none !important; }

    /* ── Design tokens ── */
    :root {
        --bg:       #05070f;
        --s1:       #0b0e1a;
        --s2:       #10141f;
        --s3:       #161b28;
        --border:   rgba(255,255,255,0.06);
        --border2:  rgba(255,255,255,0.11);
        --teal:     #00e5c8;
        --teal-dim: rgba(0,229,200,0.12);
        --purple:   #7c6df0;
        --pink:     #f05c7c;
        --amber:    #f5a623;
        --text:     #e4eaf8;
        --text2:    #9aa3be;
        --muted:    #4e576e;
        --mono:     'DM Mono', monospace;
        --display:  'Bricolage Grotesque', sans-serif;
    }

    /* ── Fixed ambient background ── */
    .lp-bg {
        position: fixed; inset: 0; z-index: 0;
        background: var(--bg);
    }
    .lp-bg::before {
        content: '';
        position: absolute; inset: -60%;
        background:
            radial-gradient(ellipse 70% 55% at 50% -10%, rgba(0,229,200,0.055) 0%, transparent 65%),
            radial-gradient(ellipse 55% 70% at 100% 60%, rgba(124,109,240,0.04) 0%, transparent 60%),
            radial-gradient(ellipse 45% 45% at 0% 80%, rgba(0,229,200,0.025) 0%, transparent 55%);
        animation: bgDrift 28s ease-in-out infinite alternate;
    }
    @keyframes bgDrift {
        0%   { transform: scale(1) rotate(0deg); }
        100% { transform: scale(1.06) rotate(1.5deg); }
    }
    /* Subtle grid */
    .lp-grid {
        position: fixed; inset: 0; z-index: 0; pointer-events: none;
        background-image:
            linear-gradient(rgba(0,229,200,0.018) 1px, transparent 1px),
            linear-gradient(90deg, rgba(0,229,200,0.018) 1px, transparent 1px);
        background-size: 72px 72px;
        mask-image: radial-gradient(ellipse 80% 80% at 50% 40%, black 20%, transparent 80%);
    }

    /* ── All content sits above bg ── */
    .lp-wrap { position: relative; z-index: 10; }

    /* ── Animated pulse dot ── */
    @keyframes pulse { 0%,100%{opacity:1;transform:scale(1)} 50%{opacity:0.3;transform:scale(0.85)} }
    .dot-live {
        display: inline-block; width: 6px; height: 6px;
        border-radius: 50%; background: var(--teal);
        animation: pulse 2s ease-in-out infinite;
        vertical-align: middle; margin-right: 7px;
    }
    @keyframes fi { to { opacity: 1; transform: translateY(0); } }
    .fi { opacity: 0; transform: translateY(18px); animation: fi 0.6s cubic-bezier(0.22,1,0.36,1) forwards; }
    .d1{animation-delay:.06s} .d2{animation-delay:.14s} .d3{animation-delay:.22s}
    .d4{animation-delay:.30s} .d5{animation-delay:.38s} .d6{animation-delay:.46s} .d7{animation-delay:.54s}

    /* ── Nav ── */
    .lp-nav {
        display: flex; align-items: center; justify-content: space-between;
        padding: 18px clamp(20px, 5vw, 60px);
        border-bottom: 1px solid var(--border);
        background: rgba(5,7,15,0.85);
        backdrop-filter: blur(24px);
        position: sticky; top: 0; z-index: 100;
    }
    .nav-brand {
        display: flex; align-items: center; gap: 10px;
        font-family: var(--display); font-weight: 800; font-size: 1.15rem;
        color: var(--text); letter-spacing: -0.03em;
    }
    .nav-logo {
        width: 32px; height: 32px; border-radius: 8px;
        background: linear-gradient(135deg, rgba(0,229,200,0.2), rgba(124,109,240,0.2));
        border: 1px solid rgba(0,229,200,0.3);
        display: flex; align-items: center; justify-content: center;
        font-size: 1rem; flex-shrink: 0;
    }
    .brand-accent { color: var(--teal); }
    .nav-right {
        display: flex; align-items: center; gap: 10px;
    }
    .nav-tag {
        font-family: var(--mono); font-size: 0.58rem;
        letter-spacing: 0.1em; text-transform: uppercase;
        padding: 5px 13px; border-radius: 100px;
        border: 1px solid var(--border); color: var(--muted);
    }
    .nav-tag-live {
        border-color: rgba(0,229,200,0.28); color: var(--teal);
        background: rgba(0,229,200,0.06);
    }
    @media (max-width: 600px) { .nav-right { display: none; } }

    /* ── Ticker ── */
    .lp-ticker {
        border-bottom: 1px solid var(--border);
        background: var(--s1);
        overflow: hidden; padding: 9px 0;
        white-space: nowrap; position: relative;
    }
    .lp-ticker::before, .lp-ticker::after {
        content: ''; position: absolute; top: 0; bottom: 0; width: 80px; z-index: 2;
    }
    .lp-ticker::before { left:0; background: linear-gradient(90deg, var(--s1), transparent); }
    .lp-ticker::after  { right:0; background: linear-gradient(-90deg, var(--s1), transparent); }
    .ticker-track {
        display: inline-flex;
        animation: tickerScroll 36s linear infinite;
    }
    @keyframes tickerScroll { 0%{transform:translateX(0)} 100%{transform:translateX(-50%)} }
    .t-item {
        font-family: var(--mono); font-size: 0.58rem;
        color: var(--muted); letter-spacing: 0.12em;
        padding: 0 32px; display: inline-flex; align-items: center; gap: 7px;
    }
    .t-dot { color: var(--teal); font-size: 0.45rem; }

    /* ══════════════════════════════════════════
       HERO — single column, centre-aligned
    ══════════════════════════════════════════ */
    .lp-hero {
        max-width: 780px;
        margin: 0 auto;
        padding: clamp(52px,8vw,96px) clamp(20px,5vw,48px) clamp(40px,6vw,72px);
        text-align: center;
        display: flex; flex-direction: column; align-items: center;
    }
    .hero-eyebrow {
        display: inline-flex; align-items: center; gap: 8px;
        font-family: var(--mono); font-size: 0.6rem;
        letter-spacing: 0.2em; text-transform: uppercase; color: var(--teal);
        padding: 6px 16px; border-radius: 100px;
        border: 1px solid rgba(0,229,200,0.22);
        background: rgba(0,229,200,0.05);
        margin-bottom: 32px;
    }
    .hero-h1 {
        font-family: var(--display); font-weight: 800;
        font-size: clamp(2.6rem, 6.5vw, 5rem);
        line-height: 1.04; letter-spacing: -0.035em;
        color: var(--text); margin: 0 0 24px 0;
    }
    .hero-h1 em { font-style: normal; color: var(--teal); }
    .hero-h1 .h1-muted {
        display: block; font-size: clamp(1.5rem,3.5vw,2.6rem);
        color: var(--muted); font-weight: 600;
        letter-spacing: -0.02em; margin-top: 8px;
    }
    .hero-sub {
        font-family: var(--mono); font-size: clamp(0.78rem,1.5vw,0.9rem);
        color: var(--text2); line-height: 1.9;
        max-width: 580px; margin: 0 0 48px 0;
    }
    .hero-sub strong { color: var(--text); }

    /* ── Hero feature pills ── */
    .hero-pills {
        display: flex; flex-wrap: wrap; justify-content: center; gap: 10px;
        margin-bottom: 48px;
    }
    .h-pill {
        display: inline-flex; align-items: center; gap: 8px;
        font-family: var(--mono); font-size: 0.62rem;
        letter-spacing: 0.06em; text-transform: uppercase;
        padding: 8px 16px; border-radius: 10px;
        border: 1px solid var(--border2);
        background: var(--s2); color: var(--text2);
        transition: border-color 0.2s, color 0.2s;
    }
    .h-pill:hover { border-color: rgba(0,229,200,0.3); color: var(--teal); }
    .h-pill-icon { font-size: 0.8rem; }

    /* ── Stats bar ── */
    .hero-stats {
        display: flex; justify-content: center; flex-wrap: wrap;
        gap: 0; border: 1px solid var(--border);
        border-radius: 14px; overflow: hidden;
        width: 100%; max-width: 560px;
        background: var(--s2);
    }
    .hs-block {
        flex: 1; min-width: 120px;
        padding: 18px 16px; text-align: center;
        border-right: 1px solid var(--border);
    }
    .hs-block:last-child { border-right: none; }
    .hs-num {
        font-family: var(--display); font-weight: 800;
        font-size: 1.5rem; color: var(--teal); line-height: 1;
    }
    .hs-label {
        font-family: var(--mono); font-size: 0.54rem;
        color: var(--muted); text-transform: uppercase; letter-spacing: 0.14em;
        margin-top: 6px;
    }

    /* ══════════════════════════════════════════
       PRICING — single column, centre-aligned
    ══════════════════════════════════════════ */
    .lp-pricing {
        max-width: 860px; margin: 0 auto;
        padding: clamp(32px,5vw,64px) clamp(20px,5vw,48px);
    }
    .section-head {
        text-align: center; margin-bottom: 40px;
    }
    .section-label {
        font-family: var(--mono); font-size: 0.58rem;
        letter-spacing: 0.22em; text-transform: uppercase; color: var(--muted);
        margin-bottom: 12px; display: block;
    }
    .section-title {
        font-family: var(--display); font-weight: 700;
        font-size: clamp(1.5rem,3vw,2rem); color: var(--text);
        letter-spacing: -0.025em; margin: 0;
    }
    .section-title em { font-style: normal; color: var(--teal); }

    /* Pricing grid — 3 cols on desktop, stacks on mobile */
    .pricing-grid {
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 16px;
    }
    @media (max-width: 720px) { .pricing-grid { grid-template-columns: 1fr; max-width: 380px; margin: 0 auto; } }

    .price-card {
        position: relative;
        background: var(--s2); border: 1px solid var(--border);
        border-radius: 18px; padding: 28px 24px 22px;
        transition: transform 0.25s, border-color 0.25s;
        overflow: hidden;
    }
    .price-card:hover { transform: translateY(-5px); border-color: var(--border2); }
    .price-card.featured {
        border-color: rgba(0,229,200,0.28);
        background: linear-gradient(155deg, rgba(0,229,200,0.035) 0%, rgba(124,109,240,0.04) 100%);
    }
    .price-card::after {
        content: ''; position: absolute; top: 0; left: 0; right: 0; height: 2px;
        opacity: 0; transition: opacity 0.3s;
        background: linear-gradient(90deg, transparent, var(--teal) 50%, transparent);
    }
    .price-card.featured::after { opacity: 1; }
    .price-card:hover::after { opacity: 1; }

    .price-badge {
        position: absolute; top: 14px; right: 14px;
        font-family: var(--mono); font-size: 0.5rem; letter-spacing: 0.12em;
        text-transform: uppercase; padding: 3px 9px; border-radius: 100px;
    }
    .badge-pop { background: rgba(0,229,200,0.1); color: var(--teal); border: 1px solid rgba(0,229,200,0.28); }
    .badge-val { background: rgba(124,109,240,0.1); color: var(--purple); border: 1px solid rgba(124,109,240,0.28); }

    .price-plan {
        font-family: var(--mono); font-size: 0.6rem;
        letter-spacing: 0.18em; text-transform: uppercase;
        color: var(--muted); margin-bottom: 14px;
    }
    .price-amount {
        font-family: var(--display); font-weight: 800; font-size: 2.6rem;
        color: var(--text); line-height: 1; letter-spacing: -0.03em;
    }
    .price-curr { font-size: 1.1rem; color: var(--muted); vertical-align: top; margin-top: 8px; display: inline-block; }
    .price-credits { font-family: var(--mono); font-size: 0.7rem; color: var(--teal); margin: 6px 0 18px; }
    .price-divider { height: 1px; background: var(--border); margin-bottom: 16px; }
    .price-features { list-style: none; padding: 0; margin: 0 0 22px; }
    .price-features li {
        font-family: var(--mono); font-size: 0.64rem; color: var(--text2);
        padding: 6px 0; display: flex; align-items: flex-start; gap: 9px;
        border-bottom: 1px solid rgba(255,255,255,0.03);
    }
    .price-features li:last-child { border-bottom: none; }
    .pf-check { color: var(--teal); flex-shrink: 0; margin-top: 1px; }

    /* ── Streamlit link-button overrides for pricing ── */
    div[data-testid="column"]:nth-child(1) .stLinkButton a,
    div[data-testid="column"]:nth-child(3) .stLinkButton a {
        background: transparent !important;
        border: 1px solid var(--border2) !important;
        color: var(--text) !important;
        font-family: var(--display) !important;
        font-weight: 700 !important; font-size: 0.72rem !important;
        border-radius: 10px !important; padding: 11px 0 !important;
        width: 100% !important; display: block !important;
        text-align: center !important; letter-spacing: 0.03em !important;
        transition: border-color 0.2s, color 0.2s !important;
    }
    div[data-testid="column"]:nth-child(1) .stLinkButton a:hover,
    div[data-testid="column"]:nth-child(3) .stLinkButton a:hover {
        border-color: rgba(0,229,200,0.4) !important; color: var(--teal) !important;
    }
    div[data-testid="column"]:nth-child(2) .stLinkButton a {
        background: linear-gradient(135deg, #00e5c8, #00c4ab) !important;
        border: none !important; color: #050b10 !important;
        font-family: var(--display) !important;
        font-weight: 800 !important; font-size: 0.72rem !important;
        border-radius: 10px !important; padding: 11px 0 !important;
        width: 100% !important; display: block !important;
        text-align: center !important; letter-spacing: 0.03em !important;
        box-shadow: 0 4px 20px rgba(0,229,200,0.32) !important;
        transition: box-shadow 0.2s, transform 0.2s !important;
    }
    div[data-testid="column"]:nth-child(2) .stLinkButton a:hover {
        box-shadow: 0 7px 28px rgba(0,229,200,0.52) !important;
        transform: translateY(-2px) !important;
    }
    .stLinkButton { margin: 0 !important; }

    /* ══════════════════════════════════════════
       ACCESS GATE — single column, centre-aligned
    ══════════════════════════════════════════ */
    .lp-gate {
        max-width: 540px; margin: 0 auto;
        padding: clamp(16px,4vw,40px) clamp(20px,5vw,48px) clamp(48px,7vw,80px);
    }
    .gate-card {
        background: var(--s2);
        border: 1px solid var(--border2);
        border-radius: 20px;
        padding: clamp(28px,5vw,44px) clamp(24px,5vw,44px);
        position: relative; overflow: hidden;
        text-align: center;
    }
    .gate-card::before {
        content: ''; position: absolute; top: 0; left: 0; right: 0; height: 1px;
        background: linear-gradient(90deg, transparent, var(--teal) 35%, var(--purple) 65%, transparent);
        opacity: 0.7;
    }
    .gate-lock {
        width: 52px; height: 52px; border-radius: 14px;
        background: linear-gradient(135deg, rgba(0,229,200,0.12), rgba(124,109,240,0.12));
        border: 1px solid rgba(0,229,200,0.2);
        display: flex; align-items: center; justify-content: center;
        font-size: 1.4rem; margin: 0 auto 20px;
    }
    .gate-title {
        font-family: var(--display); font-weight: 800; font-size: 1.5rem;
        color: var(--text); letter-spacing: -0.025em; margin-bottom: 10px;
    }
    .gate-desc {
        font-family: var(--mono); font-size: 0.7rem;
        color: var(--text2); line-height: 1.85; margin-bottom: 28px;
    }
    .gate-desc strong { color: var(--text); }
    .key-format {
        display: inline-flex; align-items: center; gap: 8px;
        font-family: var(--mono); font-size: 0.6rem; color: var(--muted);
        background: rgba(255,255,255,0.03);
        border: 1px solid var(--border); border-radius: 8px;
        padding: 7px 16px; margin-bottom: 28px;
    }
    .kf-icon { color: var(--teal); }
    /* Input + button zone */
    .gate-form { text-align: left; }
    .gate-input-label {
        font-family: var(--mono); font-size: 0.58rem;
        text-transform: uppercase; letter-spacing: 0.16em;
        color: var(--muted); margin-bottom: 8px; display: block;
    }
    .gate-links {
        margin-top: 20px; text-align: center;
        font-family: var(--mono); font-size: 0.6rem;
        color: var(--muted); line-height: 2.4;
        display: flex; flex-wrap: wrap; justify-content: center; align-items: center; gap: 4px 2px;
    }
    .gate-links a { color: var(--teal); text-decoration: none; }
    .gate-links a:hover { text-decoration: underline; text-underline-offset: 2px; }
    .gate-sep { color: var(--border2); margin: 0 6px; }
    .err-msg {
        margin-top: 12px; padding: 10px 14px;
        background: rgba(240,92,124,0.06);
        border: 1px solid rgba(240,92,124,0.22); border-radius: 10px;
        font-family: var(--mono); font-size: 0.65rem; color: var(--pink);
        display: flex; align-items: center; gap: 9px; text-align: left;
    }

    /* ── Streamlit input overrides (gate) ── */
    [data-testid="stTextInput"] > label { display: none !important; }
    [data-testid="stTextInput"] > div > div > input {
        background: rgba(255,255,255,0.035) !important;
        border: 1px solid rgba(255,255,255,0.09) !important;
        border-radius: 11px !important; color: var(--text) !important;
        font-family: var(--mono) !important; font-size: 0.82rem !important;
        padding: 14px 18px !important; letter-spacing: 0.1em !important;
        transition: border-color 0.2s, box-shadow 0.2s !important;
    }
    [data-testid="stTextInput"] > div > div > input:focus {
        border-color: rgba(0,229,200,0.45) !important;
        box-shadow: 0 0 0 3px rgba(0,229,200,0.07) !important;
        outline: none !important;
    }
    [data-testid="stTextInput"] > div > div > input::placeholder { color: #2e3650 !important; }
    [data-testid="baseButton-primary"] {
        background: linear-gradient(135deg, #00e5c8 0%, #00bfab 100%) !important;
        border: none !important; color: #04090f !important;
        font-family: var(--display) !important; font-weight: 800 !important;
        font-size: 0.85rem !important; letter-spacing: 0.04em !important;
        border-radius: 11px !important; padding: 14px !important;
        box-shadow: 0 4px 22px rgba(0,229,200,0.28) !important;
        transition: box-shadow 0.2s, transform 0.2s !important;
    }
    [data-testid="baseButton-primary"]:hover {
        box-shadow: 0 8px 32px rgba(0,229,200,0.48) !important;
        transform: translateY(-2px) !important;
    }

    /* ── Footer ── */
    .lp-footer {
        border-top: 1px solid var(--border);
        padding: 20px clamp(20px,5vw,60px);
        display: flex; flex-wrap: wrap; justify-content: space-between; align-items: center;
        gap: 10px;
        background: rgba(5,7,15,0.92);
    }
    .footer-brand {
        font-family: var(--display); font-size: 0.88rem; font-weight: 800;
        color: var(--muted); letter-spacing: -0.02em;
    }
    .footer-brand .ft { color: var(--teal); }
    .footer-copy {
        font-family: var(--mono); font-size: 0.54rem;
        color: var(--muted); letter-spacing: 0.08em; text-transform: uppercase;
        opacity: 0.5;
    }
    @media (max-width: 480px) {
        .lp-footer { flex-direction: column; text-align: center; }
        .hero-stats { flex-direction: column; }
        .hs-block { border-right: none; border-bottom: 1px solid var(--border); }
        .hs-block:last-child { border-bottom: none; }
    }
    </style>
    """, unsafe_allow_html=True)

    # ── Background layers ─────────────────────────────────────────────────────
    st.markdown('<div class="lp-bg"></div><div class="lp-grid"></div>', unsafe_allow_html=True)

    # ── Nav ───────────────────────────────────────────────────────────────────
    st.markdown("""
    <nav class="lp-nav fi d1">
        <div class="nav-brand">
            <div class="nav-logo">&#x2B21;</div>
            Panel<span class="brand-accent">Stat</span>X
        </div>
        <div class="nav-right">
            <span class="nav-tag">Econometrics Platform</span>
            <span class="nav-tag nav-tag-live"><span class="dot-live"></span>v1.0 Live</span>
        </div>
    </nav>
    """, unsafe_allow_html=True)

    # ── Ticker ────────────────────────────────────────────────────────────────
    ticker_items = [
        "Fixed Effects Estimation", "Random Effects (GLS)", "First-Difference Estimator",
        "Hausman Specification Test", "AI Explainer", "Breusch-Pagan Diagnostics",
        "Jarque-Bera Normality", "Durbin-Watson Autocorrelation", "Entity Cross-Section Plots",
        "Balanced Panel Support", "CSV & Excel Import", "DOCX Report Export",
    ]
    ticker_html = "".join(
        f'<span class="t-item"><span class="t-dot">◆</span>{item}</span>'
        for item in ticker_items
    )
    st.markdown(f"""
    <div class="lp-ticker fi d1">
        <div class="ticker-track">{ticker_html}{ticker_html}</div>
    </div>
    """, unsafe_allow_html=True)

    # ── Hero ──────────────────────────────────────────────────────────────────
    st.markdown("""
    <div class="lp-hero">

      <div class="hero-eyebrow fi d2">
        <span class="dot-live"></span>
        Advanced Econometrics Platform
      </div>

      <h1 class="hero-h1 fi d2">
        Panel Data <em>Analysis</em>
        <span class="h1-muted">Re-imagined for rigorous research.</span>
      </h1>

      <p class="hero-sub fi d3">
        Production-grade panel regression with <strong>Fixed Effects, Random Effects</strong>
        and <strong>First-Difference</strong> estimators — paired with
        <strong> AI interpretation</strong>, full diagnostic suites,
        and publication-ready DOCX reports.
      </p>

      <div class="hero-pills fi d4">
        <span class="h-pill"><span class="h-pill-icon">&#x2B21;</span> 4+ Panel Estimators</span>
        <span class="h-pill"><span class="h-pill-icon">&#x25C8;</span> AI Explainer</span>
        <span class="h-pill"><span class="h-pill-icon">&#x25C9;</span> Full Diagnostics Suite</span>
        <span class="h-pill"><span class="h-pill-icon">&#x25C6;</span> DOCX Report Export</span>
        <span class="h-pill"><span class="h-pill-icon">&#x25A3;</span> CSV &amp; Excel Import</span>
        <span class="h-pill"><span class="h-pill-icon">&#x2736;</span> Hausman Spec. Test</span>
      </div>

      <div class="hero-stats fi d5">
        <div class="hs-block">
          <div class="hs-num">4+</div>
          <div class="hs-label">Estimators</div>
        </div>
        <div class="hs-block">
          <div class="hs-num">CSV&nbsp;&amp;&nbsp;XLS</div>
          <div class="hs-label">Data Formats</div>
        </div>
        <div class="hs-block">
          <div class="hs-num">AI Model</div>
          <div class="hs-label">Intelligent layer</div>
        </div>
        <div class="hs-block">
          <div class="hs-num">DOCX</div>
          <div class="hs-label">Report Export</div>
        </div>
      </div>

    </div>
    """, unsafe_allow_html=True)



    # ── Pricing header ─────────────────────────────────────────────────────
    st.markdown("""
    <div class="lp-pricing">
      <div class="section-head fi d5">
        <span class="section-label">Pay-As-You-Go Credits &middot; No Subscription</span>
        <p class="section-title">Simple, <em>transparent</em> pricing</p>
      </div>
    </div>
    """, unsafe_allow_html=True)

        # ── Pricing cards ─────────────────────────────────────────────────────
    pc1, pc2, pc3 = st.columns(3, gap="small")

    with pc1:
        st.markdown("""
        <div class="price-card fi d3">
          <div class="price-plan">Starter</div>
          <div class="price-amount"><span class="price-curr">$</span>10</div>
          <div class="price-credits">5 Analysis Credits</div>
          <div class="price-divider"></div>
          <ul class="price-features">
            <li><span class="pf-check">&#x2713;</span>5 full regression runs</li>
            <li><span class="pf-check">&#x2713;</span>All 4 estimators</li>
            <li><span class="pf-check">&#x2713;</span>AI explainer included</li>
            <li><span class="pf-check">&#x2713;</span>DOCX report export</li>
          </ul>
        </div>
        """, unsafe_allow_html=True)
        st.link_button("Get Started →", "https://flutterwave.com/pay/j5m67hrqr4iq",
                            width='stretch')

    with pc2:
        st.markdown("""
        <div class="price-card featured fi d4">
          <span class="price-badge badge-pop">Most Popular</span>
          <div class="price-plan">Standard</div>
          <div class="price-amount"><span class="price-curr">$</span>25</div>
          <div class="price-credits">20 Analysis Credits</div>
          <div class="price-divider"></div>
          <ul class="price-features">
            <li><span class="pf-check">&#x2713;</span>20 full regression runs</li>
            <li><span class="pf-check">&#x2713;</span>All 4 estimators</li>
            <li><span class="pf-check">&#x2713;</span>AI explainer included</li>
            <li><span class="pf-check">&#x2713;</span>DOCX report export</li>
          </ul>
        </div>
        """, unsafe_allow_html=True)
        st.link_button("Buy Credits →", "https://flutterwave.com/pay/txyljcuqfsel",
                           width='stretch')

    with pc3:
        st.markdown("""
        <div class="price-card fi d5">
          <span class="price-badge badge-val">Best Value</span>
          <div class="price-plan">Team</div>
          <div class="price-amount"><span class="price-curr">$</span>100</div>
          <div class="price-credits">100 Analysis Credits</div>
          <div class="price-divider"></div>
          <ul class="price-features">
            <li><span class="pf-check">&#x2713;</span>100 full regression runs</li>
            <li><span class="pf-check">&#x2713;</span>All 4 estimators</li>
            <li><span class="pf-check">&#x2713;</span>AI explainer included</li>
            <li><span class="pf-check">&#x2713;</span>DOCX report export</li>
          </ul>
        </div>
        """, unsafe_allow_html=True)
        st.link_button("Buy Credits →", "https://flutterwave.com/pay/yw7k3gmyjoud",
                           width='stretch')

    # ── Access Gate ───────────────────────────────────────────────────────────
    st.markdown("""
    <div class="lp-gate fi d6">
      <div class="gate-card">
        <div class="gate-lock">&#x1F511;</div>
        <div class="gate-title">Enter Your Access Key</div>
        <div class="gate-desc">
          PanelStatX uses a <strong>credit-based, pay-as-you-go model</strong>.
          Each analysis run deducts one credit from your unique,
          <strong>non-expiring</strong> access key.
          No recurring billing. No monthly subscriptions.
        </div>
        <div class="key-format">
          <span class="kf-icon">&#x25C8;</span>
          Key format: PSX-XXXX-XXXX-XXXX
        </div>
        <div class="gate-form">
          <span class="gate-input-label">Access Key</span>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Inject Streamlit widgets into gate card ───────────────────────────────
    st.markdown("""
    <style>
    [data-testid="stVerticalBlock"] > [data-testid="stVerticalBlock"] { margin-top: -18px; }
    </style>
    """, unsafe_allow_html=True)

    _, gate_input_col, _ = st.columns([1, 2, 1])
    with gate_input_col:
        entered_key = st.text_input(
            "Access Key",
            type="password",
            placeholder="PSX-XXXX-XXXX-XXXX",
            label_visibility="collapsed",
        )
        unlock_btn = st.button("⬡  Unlock PanelStatX", width='stretch', type="primary")

        if st.session_state.access_error:
            st.markdown(f"""
            <div class="err-msg">
                <span>&#x2715;</span> {st.session_state.access_error}
            </div>
            """, unsafe_allow_html=True)
            
        st.markdown("""
        <div class="gate-links">
            <a href="https://wa.me/2348096506034" target="_blank" rel="noopener noreferrer">&#x1F464; Get Access Key</a>
            <span class="gate-sep">&nbsp;|&nbsp;</span>
            <a href="https://app.box.com/s/vw4c6u10bv0z8ngarzj73ej18t74e3wl" target="_blank" rel="noopener noreferrer">&#x1F4CB; User Guide</a>
            <span class="gate-sep">&nbsp;|&nbsp;</span>
            <a href="mailto:Abdulwrite77@gmail.com">&#x2699;&#xFE0F; Support</a>
        </div>
        """, unsafe_allow_html=True)
 

    # ── Footer ────────────────────────────────────────────────────────────────
    st.markdown("""
    <div class="lp-footer fi d7">
        <div class="footer-brand">Panel<span class="ft">Stat</span>X</div>
        <div class="footer-copy">Panel Regression Analysis &middot; Powered by AI Model &middot; Credit-based access</div>
    </div>
    """, unsafe_allow_html=True)

    # ── Unlock logic ──────────────────────────────────────────────────────────
    if unlock_btn:
        if not entered_key:
            st.session_state.access_error = "Please enter your access key."
            st.rerun()
        else:
            with st.spinner("Verifying key…"):
                record = lookup_key(entered_key)
            if record is None:
                st.session_state.access_error = "Invalid access key. Please try again."
                st.rerun()
            elif record["credits"] <= 0:
                st.session_state.access_error = (
                    "Your credits have been exhausted. "
                    "Please purchase more credits to continue."
                )
                st.rerun()
            else:
                st.session_state.access_granted = True
                st.session_state.access_error   = ""
                st.session_state.user_key        = record["key"]
                st.session_state.user_credits    = record["credits"]
                st.session_state.user_email      = record["email"]
                st.session_state.user_row        = record["row_index"]
                st.rerun()

    st.stop()



# ═══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════════════════

with st.sidebar:
    st.markdown("""
    <div style="padding:16px 0 24px 0; border-bottom:1px solid var(--border); margin-bottom:20px;">
        <div style="font-family:'Syne',sans-serif;font-size:1.4rem;font-weight:800;color:var(--text);letter-spacing:-0.02em;">
            &#x2B21; Panel<span style="color:var(--accent);">Stat</span>X
        </div>
        <div style="font-family:'DM Mono',monospace;font-size:0.68rem;color:var(--muted);margin-top:4px;letter-spacing:0.08em;">
            PANEL REGRESSION ENGINE v1.0
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Credit HUD ────────────────────────────────────────────────────────────
    credits_left = st.session_state.user_credits
    credit_color = "#00e5c8" if credits_left > 5 else "#f5a623" if credits_left > 1 else "#f05c7c"
    credit_label = "Credits remaining" if credits_left > 1 else ("1 credit left!" if credits_left == 1 else "No credits")
    email_display = st.session_state.user_email or "—"
    st.markdown(f"""
    <div style="background:var(--surface2);border:1px solid var(--border);border-radius:8px;
                padding:14px 16px;margin-bottom:20px;font-family:'DM Mono',monospace;">
        <div style="font-size:0.6rem;text-transform:uppercase;letter-spacing:0.12em;
                    color:var(--muted);margin-bottom:8px;">Account</div>
        <div style="font-size:0.7rem;color:var(--muted);margin-bottom:10px;
                    overflow:hidden;text-overflow:ellipsis;white-space:nowrap;">{email_display}</div>
        <div style="display:flex;align-items:baseline;gap:6px;">
            <span style="font-family:'Syne',sans-serif;font-size:1.6rem;font-weight:700;
                         color:{credit_color};line-height:1;">{credits_left}</span>
            <span style="font-size:0.62rem;color:var(--muted);text-transform:uppercase;
                         letter-spacing:0.1em;">{credit_label}</span>
        </div>
        <div style="margin-top:8px;background:var(--border);border-radius:2px;height:3px;">
            <div style="height:3px;border-radius:2px;background:{credit_color};
                        width:{min(100, credits_left * 10)}%;transition:width 0.4s;"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if credits_left <= 0:
        st.error("⚠ No credits remaining. Please purchase more to run analyses.")
        st.markdown("""
        <div style="text-align:center;font-family:'DM Mono',monospace;font-size:0.7rem;
                    color:var(--muted);padding:8px 0;">
            <a href="#" style="color:var(--accent);text-decoration:none;">Buy more credits →</a>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<div style="font-size:0.7rem;text-transform:uppercase;letter-spacing:0.12em;color:var(--muted);margin-bottom:8px;">Data Source</div>', unsafe_allow_html=True)
    data_src = st.radio("", ["Use Demo Dataset", "Upload File"], label_visibility="collapsed")

    if data_src == "Upload File":
        uploaded = st.file_uploader(
            "Upload panel data",
            type=["csv", "xlsx", "xls"],
            label_visibility="collapsed",
            help="Accepts CSV (.csv) or Excel (.xlsx / .xls) files",
        )
        if uploaded:
            try:
                fname = uploaded.name.lower()
                if fname.endswith(".csv"):
                    st.session_state.df = pd.read_csv(uploaded)
                else:
                    # Excel: let user pick sheet if multiple exist
                    xf = pd.ExcelFile(uploaded)
                    if len(xf.sheet_names) > 1:
                        sheet = st.selectbox("Sheet", xf.sheet_names)
                    else:
                        sheet = xf.sheet_names[0]
                    st.session_state.df = pd.read_excel(uploaded, sheet_name=sheet)
                st.success(f"Loaded {st.session_state.df.shape[0]:,} rows × {st.session_state.df.shape[1]} cols")
            except Exception as e:
                st.error(f"Could not read file: {e}")
    else:
        if st.button("Load Demo Data", use_container_width=True):
            st.session_state.df = generate_demo_panel()
            st.success("Demo panel loaded!")

    st.markdown("---")

    if st.session_state.df is not None:
        df = st.session_state.df
        cols = df.columns.tolist()

        st.markdown('<div style="font-size:0.7rem;text-transform:uppercase;letter-spacing:0.12em;color:var(--muted);margin-bottom:8px;">Variable Mapping</div>', unsafe_allow_html=True)
        entity_col = st.selectbox("Entity / Panel ID", cols, index=cols.index("entity") if "entity" in cols else 0)
        time_col   = st.selectbox("Time Variable", cols, index=cols.index("year") if "year" in cols else 1)
        y_col      = st.selectbox("Dependent Variable (Y)", [c for c in cols if c not in [entity_col, time_col]],
                                   index=0)
        x_candidates = [c for c in cols if c not in [entity_col, time_col, y_col]]
        x_cols = st.multiselect("Independent Variables (X)", x_candidates, default=x_candidates[:3] if len(x_candidates) >= 3 else x_candidates)

        st.markdown("---")
        st.markdown('<div style="font-size:0.7rem;text-transform:uppercase;letter-spacing:0.12em;color:var(--muted);margin-bottom:8px;">Estimator</div>', unsafe_allow_html=True)
        model_type = st.selectbox("", ["Fixed Effects (Two-Way)", "Fixed Effects (Entity)", "Random Effects (GLS)", "First Difference", "Pooled OLS"], label_visibility="collapsed")
        st.session_state.model_type = model_type

        st.markdown("---")
        run_btn = st.button(
            "Run Analysis",
            use_container_width=True,
            type="primary",
            disabled=(st.session_state.user_credits <= 0),
        )

        # ── Analysis status feedback label ────────────────────────────────────
        if st.session_state.results is not None:
            st.markdown("""
            <div style="
                display:flex; align-items:center; gap:8px;
                margin-top:8px; padding:9px 14px;
                background:rgba(34,211,160,0.07);
                border:1px solid rgba(34,211,160,0.25);
                border-radius:8px;
                font-family:'DM Mono',monospace;
                font-size:0.68rem; color:#22d3a0;
                letter-spacing:0.04em;
            ">
                <span style="font-size:0.8rem;">&#x2713;</span>
                Analysis complete
            </div>
            """, unsafe_allow_html=True)
        elif st.session_state.df is not None and not (st.session_state.user_credits <= 0):
            st.markdown("""
            <div style="
                display:flex; align-items:center; gap:8px;
                margin-top:8px; padding:9px 14px;
                background:rgba(255,255,255,0.02);
                border:1px solid rgba(255,255,255,0.06);
                border-radius:8px;
                font-family:'DM Mono',monospace;
                font-size:0.68rem; color:#4e576e;
                letter-spacing:0.04em;
            ">
                <span style="font-size:0.8rem;">&#x25CB;</span>
                Ready to analyse
            </div>
            """, unsafe_allow_html=True)
    else:
        run_btn = False
        entity_col = time_col = y_col = "—"
        x_cols = []
        model_type = "Fixed Effects (Two-Way)"

    # ── New Analysis reset button ───────────────────────────────────────────────
    st.markdown("---")
    if st.button("New Analysis", use_container_width=True):
        st.session_state.df = None
        st.session_state.results = None
        st.session_state.ai_explanation = ""
        st.rerun()


# ═══════════════════════════════════════════════════════════════════════════════
# RUN MODEL
# ═══════════════════════════════════════════════════════════════════════════════

if run_btn and st.session_state.df is not None and x_cols:
    # ── Credit guard ──────────────────────────────────────────────────────────
    if st.session_state.user_credits <= 0:
        st.error("⚠ You have no credits remaining. Please purchase more to run analyses.")
        st.stop()

    df = st.session_state.df
    with st.spinner("Running regression…"):
        try:
            if model_type == "Pooled OLS":
                result_df, resid, y_hat, stats, vcov = run_ols(df, y_col, x_cols)
            elif model_type == "First Difference":
                result_df, resid, y_hat, stats, vcov = run_fd(df, y_col, x_cols, entity_col, time_col)
            elif model_type == "Fixed Effects (Entity)":
                result_df, resid, y_hat, stats, vcov = run_within(df, y_col, x_cols, entity_col, time_col)
            elif model_type == "Random Effects (GLS)":
                result_df, resid, y_hat, stats, vcov = run_re(df, y_col, x_cols, entity_col, time_col)
            else:  # Fixed Effects (Two-Way)
                result_df, resid, y_hat, stats, vcov = run_within(df, y_col, x_cols, entity_col, time_col)

            # ── Breusch-Pagan heteroskedasticity test ─────────────────────────
            # Build regressor matrix for BP (with intercept)
            try:
                _bp_X = np.column_stack([np.ones(len(resid))] +
                                         [df[c].values[:len(resid)] for c in x_cols])
                bp_stat, bp_p, bp_k = breusch_pagan_test(resid, _bp_X)
                stats["BP_stat"] = bp_stat
                stats["BP_p"]    = bp_p
            except Exception:
                stats["BP_stat"] = None
                stats["BP_p"]    = None

            # ── Hausman test (RE vs FE) ────────────────────────────────────────
            hausman_result = None
            if model_type == "Random Effects (GLS)":
                try:
                    _, _, _, _, fe_vcov = run_within(df, y_col, x_cols, entity_col, time_col)
                    fe_res, _, _, _, _  = run_within(df, y_col, x_cols, entity_col, time_col)
                    fe_coef = fe_res["Coeff"].values
                    re_coef = result_df[result_df["Variable"] != "const"]["Coeff"].values
                    # align – both should be len(x_cols)
                    h_stat, h_p, h_df = hausman_test(fe_coef, re_coef, fe_vcov, vcov[1:, 1:])
                    hausman_result = {"stat": h_stat, "p": h_p, "df": h_df}
                except Exception:
                    pass

            st.session_state.results = {
                "result_df": result_df, "resid": resid, "y_hat": y_hat,
                "stats": stats, "y_col": y_col, "x_cols": x_cols,
                "entity_col": entity_col, "time_col": time_col,
                "hausman": hausman_result,
            }
            st.session_state.ai_explanation = ""

            # ── Deduct 1 credit ───────────────────────────────────────────────
            new_credits = deduct_credit(
                st.session_state.user_row,
                st.session_state.user_credits,
            )
            st.session_state.user_credits = new_credits
            if new_credits == 0:
                st.warning("⚠ You just used your last credit. Purchase more to run further analyses.")
            elif new_credits <= 3:
                st.warning(f"⚠ Only {new_credits} credit(s) remaining.")

        except Exception as e:
            st.error(f"Regression error: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN LAYOUT
# ═══════════════════════════════════════════════════════════════════════════════

# Hero
st.markdown("""
<div class="hero">
    <div class="hero-title">Panel<span>Stat</span>X</div>
    <div class="hero-sub">⬡ Panel Regression Analysis System · AI-Powered Econometrics</div>
</div>
""", unsafe_allow_html=True)

if st.session_state.df is None:
    # ── Landing state ──────────────────────────────────────────────────────────
    col1, col2, col3 = st.columns(3)
    for col, icon, title, desc in [
        (col1, "⬡", "Panel-Ready", "Fixed effects, first-difference, and pooled OLS estimators built for longitudinal data."),
        (col2, "◈", "Diagnostic Suite", "Residual analysis, heteroskedasticity checks, Hausman test, and entity plots."),
        (col3, "⬟", "AI Explainer", "AI model interprets your regression output in plain language — coefficients, fit, and caveats."),
    ]:
        with col:
            st.markdown(f"""
            <div class="scard" style="text-align:center;padding:32px 20px;">
                <div style="font-size:2rem;margin-bottom:12px;color:var(--accent);">{icon}</div>
                <div style="font-family:'Syne',sans-serif;font-size:1rem;font-weight:700;color:var(--text);margin-bottom:8px;">{title}</div>
                <div style="font-size:0.78rem;color:var(--muted);line-height:1.6;">{desc}</div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("""
    <div style="text-align:center;margin-top:40px;padding:32px;background:var(--surface2);border:1px dashed var(--border);border-radius:8px;">
        <div style="font-family:'Syne',sans-serif;font-size:1rem;color:var(--muted);">
            ← Load demo data or upload a CSV from the sidebar to begin
        </div>
    </div>
    """, unsafe_allow_html=True)
    st.stop()


# ── Data has been loaded ────────────────────────────────────────────────────
df = st.session_state.df
res = st.session_state.results

# ── Quick dataset stats bar ───────────────────────────────────────────────────
n_e = df[entity_col].nunique() if entity_col in df.columns else "—"
n_t = df[time_col].nunique() if time_col in df.columns else "—"
st.markdown(f"""
<div style="margin-bottom:24px;">
    <span class="stat-pill">Entities <b>{n_e}</b></span>
    <span class="stat-pill">Periods <b>{n_t}</b></span>
    <span class="stat-pill">Observations <b>{len(df):,}</b></span>
    <span class="stat-pill">Estimator <b>{st.session_state.model_type}</b></span>
    <span class="badge badge-teal" style="margin-left:4px;">READY</span>
</div>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════════
# TABS
# ═══════════════════════════════════════════════════════════════════════════════

tab_data, tab_results, tab_diagnostics, tab_entity, tab_ai = st.tabs([
    "⬡ Data Explorer", "◈ Results", "⬟ Diagnostics", "⬢ Entity Plots", "✦ AI Explainer"
])


# ──────────────────────────────────────────────────────────────────────────────
# TAB 1 · DATA EXPLORER
# ──────────────────────────────────────────────────────────────────────────────

with tab_data:
    c1, c2 = st.columns([2, 1])
    with c1:
        st.markdown('<div class="scard-title">Dataset Preview</div>', unsafe_allow_html=True)
        st.dataframe(df.head(100), width='stretch', height=320)
    with c2:
        st.markdown('<div class="scard-title">Summary Statistics</div>', unsafe_allow_html=True)
        st.dataframe(df.describe().round(3), width='stretch', height=320)

    st.markdown("---")
    st.markdown('<div class="scard-title">Correlation Heatmap</div>', unsafe_allow_html=True)
    num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    if len(num_cols) >= 2:
        corr = df[num_cols].corr().round(3)
        fig_corr = go.Figure(go.Heatmap(
            z=corr.values, x=corr.columns, y=corr.index,
            colorscale=[[0, "#f05c7c"], [0.5, "#111318"], [1, "#00e5c8"]],
            zmid=0, text=corr.values.round(2),
            texttemplate="%{text}", showscale=True,
        ))
        fig_corr.update_layout(title="Pearson Correlation Matrix", height=380, **PLOTLY_THEME)
        st.plotly_chart(fig_corr,  width='stretch')

    # Distribution of Y
    if y_col in df.columns:
        st.markdown('<div class="scard-title" style="margin-top:16px;">Dependent Variable Distribution</div>', unsafe_allow_html=True)
        fig_dist = px.histogram(df, x=y_col, nbins=40, color_discrete_sequence=["#00e5c8"])
        fig_dist.update_layout(title=f"Distribution of {y_col}", height=300, bargap=0.05, **PLOTLY_THEME)
        st.plotly_chart(fig_dist,  width='stretch')


# ──────────────────────────────────────────────────────────────────────────────
# TAB 2 · RESULTS
# ──────────────────────────────────────────────────────────────────────────────

with tab_results:
    if res is None:
        st.info("Run the analysis from the sidebar to view regression results.")
    else:
        result_df = res["result_df"]
        stats = res["stats"]

        # ── Model fit summary ──────────────────────────────────────────────────
        st.markdown('<div class="scard-title">Model Fit</div>', unsafe_allow_html=True)
        mc = st.columns(8)
        f_label = f"{stats.get('F_stat', 0):.3f}" if stats.get('F_stat') is not None else "—"
        fp_label = f"{stats.get('F_p', 1):.4f}" if stats.get('F_p') is not None else "—"
        for col, label, val in [
            (mc[0], "R²",        f"{stats['R2']:.4f}"),
            (mc[1], "Adj. R²",   f"{stats['R2_adj']:.4f}"),
            (mc[2], "N",         f"{stats['N']:,}"),
            (mc[3], "Variables", f"{stats['k']}"),
            (mc[4], "AIC",       f"{stats['AIC']:.2f}"),
            (mc[5], "BIC",       f"{stats['BIC']:.2f}"),
            (mc[6], "F-stat",    f_label),
            (mc[7], "F p-value", fp_label),
        ]:
            with col:
                st.metric(label, val)

        # F-stat interpretation
        if stats.get("F_p") is not None:
            if stats["F_p"] < 0.05:
                st.success(f"✓ F-statistic ({stats['F_stat']:.3f}) is significant (p={stats['F_p']:.4f}) — regressors jointly explain the outcome.")
            else:
                st.warning(f"⚠ F-statistic ({stats['F_stat']:.3f}) is not significant (p={stats['F_p']:.4f}) — regressors may not jointly explain the outcome.")

        # RE variance components
        if st.session_state.model_type == "Random Effects (GLS)":
            st.markdown("---")
            st.markdown('<div class="scard-title">Random Effects Variance Components</div>', unsafe_allow_html=True)
            rc1, rc2, rc3 = st.columns(3)
            with rc1: st.metric("σ²ᵤ (between)", f"{stats.get('sigma_u2', 0):.6f}")
            with rc2: st.metric("σ²ₑ (within)",  f"{stats.get('sigma_e2', 0):.6f}")
            with rc3: st.metric("θ (GLS weight)", f"{stats.get('theta', 0):.4f}")
            st.caption("θ → 1 means FE dominates; θ → 0 means OLS/pooled dominates.")

        # Hausman test
        hausman_res = res.get("hausman")
        if hausman_res and hausman_res.get("stat") is not None:
            st.markdown("---")
            st.markdown('<div class="scard-title">Hausman Specification Test (RE vs FE)</div>', unsafe_allow_html=True)
            hc1, hc2, hc3 = st.columns(3)
            with hc1: st.metric("χ² statistic", f"{hausman_res['stat']:.4f}")
            with hc2: st.metric("p-value",       f"{hausman_res['p']:.4f}")
            with hc3: st.metric("df",             f"{hausman_res['df']}")
            if hausman_res["p"] < 0.05:
                st.warning("⚠ Hausman test rejects RE (p < 0.05) — Fixed Effects estimator is preferred (endogenous individual effects suspected).")
            else:
                st.success("✓ Hausman test does not reject RE (p ≥ 0.05) — Random Effects estimator is consistent and efficient.")

        st.markdown("---")

        # ── Coefficient table ──────────────────────────────────────────────────
        st.markdown('<div class="scard-title">Coefficient Estimates</div>', unsafe_allow_html=True)
        display = result_df.copy()
        display["Stars"]  = display["p_value"].apply(significance_stars)
        display["Sig"]    = display["p_value"].apply(
            lambda p: "✓ Significant" if p < 0.05 else "✗ Not sig.")
        display = display.rename(columns={
            "Variable": "Variable", "Coeff": "Coeff.",
            "Std_Err": "Std. Err.", "t_stat": "t-stat", "p_value": "p-value"
        })
        st.dataframe(
            display.style
                .format({"Coeff.": "{:.4f}", "Std. Err.": "{:.4f}",
                         "t-stat": "{:.3f}", "p-value": "{:.4f}"})
                .map(lambda v: "color: #00e5c8" if v == "✓ Significant" else "color: #6b7a9a", subset=["Sig"]),
            use_container_width=True, hide_index=True
        )
        st.caption("*p<0.1  **p<0.05  ***p<0.01")

        st.markdown("---")

        # ── Coefficient plot ───────────────────────────────────────────────────
        st.markdown('<div class="scard-title">Coefficient Plot (with 95% CI)</div>', unsafe_allow_html=True)
        rd = res["result_df"]
        ci_lo = rd["Coeff"] - 1.96 * rd["Std_Err"]
        ci_hi = rd["Coeff"] + 1.96 * rd["Std_Err"]
        colors = ["#00e5c8" if p < 0.05 else "#6b7a9a" for p in rd["p_value"]]

        fig_coef = go.Figure()
        fig_coef.add_hline(y=0, line_dash="dash", line_color="#1f2535")
        for pos, (i, row) in enumerate(rd.iterrows()):
            lo, hi = ci_lo.iloc[pos], ci_hi.iloc[pos]
            fig_coef.add_trace(go.Scatter(
                x=[lo, hi], y=[row["Variable"], row["Variable"]],
                mode="lines", line=dict(color="#1f2535", width=2),
                showlegend=False
            ))
        fig_coef.add_trace(go.Scatter(
            x=rd["Coeff"], y=rd["Variable"], mode="markers",
            marker=dict(size=10, color=colors, line=dict(width=1, color="#0a0c10")),
            name="Coefficient", showlegend=False
        ))
        fig_coef.update_layout(title="Coefficients with 95% Confidence Intervals",
                                height=max(280, len(rd) * 55), **PLOTLY_THEME)
        st.plotly_chart(fig_coef,  width='stretch')

        # ── Download Report ────────────────────────────────────────────────────
        st.markdown("---")
        st.markdown('<div class="scard-title">Export Report</div>', unsafe_allow_html=True)
        st.markdown("""
        <div style="font-family:'DM Mono',monospace;font-size:0.76rem;color:var(--muted);margin-bottom:12px;line-height:1.7;">
            Downloads a fully formatted MS Word (.docx) report containing: model fit statistics,
            coefficient estimates table with significance stars &amp; 95% CI, residual diagnostics table,
            and the AI write-up if already generated in the AI Explainer tab.
        </div>
        """, unsafe_allow_html=True)

        dl_col1, dl_col2 = st.columns([1, 2])
        with dl_col1:
            try:
                import datetime
                docx_bytes = build_docx_report(
                    res,
                    st.session_state.model_type,
                    ai_explanation=st.session_state.get("ai_explanation", ""),
                )
                fname = f"PanelStatX_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                st.download_button(
                    label="⬇  Download Report (.docx)",
                    data=docx_bytes,
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary",
                )
            except Exception as e:
                st.error(f"Report generation error: {e}")
        with dl_col2:
            st.markdown("""
            <div style="font-family:'DM Mono',monospace;font-size:0.7rem;color:var(--muted);padding:10px 0;line-height:1.8;">
                ✦ Tip: Generate the AI Explanation first (AI Explainer tab),
                then return here to download the complete report including the write-up.
            </div>
            """, unsafe_allow_html=True)


# ──────────────────────────────────────────────────────────────────────────────
# TAB 3 · DIAGNOSTICS
# ──────────────────────────────────────────────────────────────────────────────

with tab_diagnostics:
    if res is None:
        st.info("Run the analysis first.")
    else:
        resid = res["resid"]
        y_hat = res["y_hat"]

        dc1, dc2 = st.columns(2)

        # Residuals vs Fitted
        with dc1:
            fig_rv = go.Figure()
            fig_rv.add_hline(y=0, line_dash="dash", line_color="#f05c7c", line_width=1)
            fig_rv.add_trace(go.Scatter(
                x=y_hat, y=resid, mode="markers",
                marker=dict(size=5, color="#00e5c8", opacity=0.6),
                name="Residual"
            ))
            fig_rv.update_layout(title="Residuals vs Fitted",
                                   xaxis_title="Fitted", yaxis_title="Residual", height=340, **PLOTLY_THEME)
            st.plotly_chart(fig_rv, width='stretch')

        # Q-Q Plot
        with dc2:
            from scipy import stats as sc_stats
            resid_clean = np.asarray(resid, dtype=float)
            resid_clean = resid_clean[np.isfinite(resid_clean)]
            probplot_result = sc_stats.probplot(resid_clean)
            osm, osr = probplot_result[0]
            slope, intercept, _ = probplot_result[1]
            fig_qq = go.Figure()
            fig_qq.add_trace(go.Scatter(x=list(osm), y=list(osr), mode="markers",
                                         marker=dict(size=4, color="#7c6df0", opacity=0.7), name="Residuals"))
            fig_qq.add_trace(go.Scatter(x=[float(min(osm)), float(max(osm))],
                                         y=[slope * float(min(osm)) + intercept, slope * float(max(osm)) + intercept],
                                         mode="lines", line=dict(color="#f05c7c", dash="dash"), name="Normal"))
            fig_qq.update_layout(title="Normal Q-Q Plot",
                                   xaxis_title="Theoretical Quantiles",
                                   yaxis_title="Sample Quantiles", height=340, **PLOTLY_THEME)
            st.plotly_chart(fig_qq, width='stretch')

        dc3, dc4 = st.columns(2)

        # Residual distribution
        with dc3:
            fig_rh = px.histogram(x=resid, nbins=40, color_discrete_sequence=["#7c6df0"])
            fig_rh.update_layout(title="Residual Distribution", xaxis_title="Residual",
                                   height=300, bargap=0.05, **PLOTLY_THEME)
            st.plotly_chart(fig_rh, width='stretch')

        # Scale-Location
        with dc4:
            fig_sl = go.Figure()
            fig_sl.add_trace(go.Scatter(
                x=y_hat, y=np.sqrt(np.abs(resid)), mode="markers",
                marker=dict(size=5, color="#f5a623", opacity=0.6)
            ))
            fig_sl.update_layout(title="Scale-Location (√|Residual| vs Fitted)",
                                   xaxis_title="Fitted", yaxis_title="√|Residual|", height=300, **PLOTLY_THEME)
            st.plotly_chart(fig_sl, width='stretch')

        # Residual stats
        st.markdown("---")
        st.markdown('<div class="scard-title">Residual Diagnostics Summary</div>', unsafe_allow_html=True)
        from scipy import stats as sc_stats
        resid_arr = np.asarray(resid, dtype=float)
        resid_arr = resid_arr[np.isfinite(resid_arr)]
        jb_stat, jb_p = sc_stats.jarque_bera(resid_arr)
        dw = np.sum(np.diff(resid_arr)**2) / max(np.sum(resid_arr**2), 1e-10)

        dc5, dc6, dc7, dc8 = st.columns(4)
        with dc5: st.metric("Mean Residual", f"{np.mean(resid_arr):.4f}")
        with dc6: st.metric("Std Residual", f"{np.std(resid_arr):.4f}")
        with dc7: st.metric("Jarque-Bera p", f"{jb_p:.4f}")
        with dc8: st.metric("Durbin-Watson", f"{dw:.4f}")

        if jb_p < 0.05:
            st.warning("⚠ Jarque-Bera test rejects normality (p < 0.05). Consider robust standard errors.")
        if dw < 1.5 or dw > 2.5:
            st.warning(f"⚠ Durbin-Watson = {dw:.3f} suggests possible autocorrelation.")
        else:
            st.success("✓ Durbin-Watson statistic is in the acceptable range.")

        # ── Breusch-Pagan Heteroskedasticity Test ─────────────────────────────
        st.markdown("---")
        st.markdown('<div class="scard-title">Breusch-Pagan Test for Heteroskedasticity</div>', unsafe_allow_html=True)
        bp_stat = res["stats"].get("BP_stat")
        bp_p    = res["stats"].get("BP_p")
        if bp_stat is not None and bp_p is not None:
            bpc1, bpc2 = st.columns(2)
            with bpc1: st.metric("BP LM Statistic", f"{bp_stat:.4f}")
            with bpc2: st.metric("BP p-value",       f"{bp_p:.4f}")
            if bp_p < 0.05:
                st.warning("⚠ Breusch-Pagan test rejects homoskedasticity (p < 0.05). "
                           "Heteroskedastic errors detected — consider heteroskedasticity-robust (HC) standard errors.")
            else:
                st.success("✓ Breusch-Pagan test does not reject homoskedasticity (p ≥ 0.05).")
            st.caption("H₀: Constant variance (homoskedasticity). LM ~ χ²(k).")
        else:
            st.info("Breusch-Pagan test could not be computed for this estimator/data combination.")


# ──────────────────────────────────────────────────────────────────────────────
# TAB 4 · ENTITY PLOTS
# ──────────────────────────────────────────────────────────────────────────────

with tab_entity:
    if entity_col not in df.columns or time_col not in df.columns:
        st.info("Entity and time columns not set.")
    else:
        y_plot = y_col if y_col in df.columns else df.select_dtypes(np.number).columns[0]

        ec1, ec2 = st.columns([1, 3])
        with ec1:
            entities_avail = sorted(df[entity_col].unique())
            selected_entities = st.multiselect(
                "Select entities to plot",
                entities_avail,
                default=entities_avail[:6] if len(entities_avail) >= 6 else entities_avail
            )
        with ec2:
            x_axis = st.selectbox("X axis", [time_col] + [c for c in df.columns if c not in [entity_col]], index=0)

        if selected_entities:
            plot_df = df[df[entity_col].isin(selected_entities)]

            fig_ep = px.line(
                plot_df, x=x_axis, y=y_plot, color=entity_col,
                markers=True,
                color_discrete_sequence=["#00e5c8", "#7c6df0", "#f05c7c", "#f5a623",
                                          "#22d3a0", "#60a5fa", "#fb923c", "#a78bfa"],
            )
            fig_ep.update_layout(title=f"{y_plot} over {x_axis} by {entity_col}",
                                   height=440, **PLOTLY_THEME)
            st.plotly_chart(fig_ep, width='stretch')

            # Entity mean bar
            means = df.groupby(entity_col)[y_plot].mean().sort_values(ascending=False)
            fig_bar = px.bar(
                x=means.index, y=means.values,
                color=means.values,
                color_continuous_scale=["#111318", "#00e5c8"],
                labels={"x": entity_col, "y": f"Mean {y_plot}"},
            )
            fig_bar.update_layout(title=f"Entity Mean of {y_plot}", height=320,
                                   coloraxis_showscale=False, **PLOTLY_THEME)
            st.plotly_chart(fig_bar, width='stretch')


# ──────────────────────────────────────────────────────────────────────────────
# TAB 5 · AI EXPLAINER
# ──────────────────────────────────────────────────────────────────────────────

with tab_ai:
    st.markdown('<div class="scard-title">AI Regression Explainer</div>', unsafe_allow_html=True)

    if res is None:
        st.info("Run the analysis first to unlock the AI explainer.")
    else:
        result_df = res["result_df"]
        stats = res["stats"]

        # Build a rich summary for GPT-4o Model
        coeff_table = result_df.to_string(index=False)
        hausman_res = res.get("hausman")
        hausman_str = ""
        if hausman_res and hausman_res.get("stat") is not None:
            hausman_str = f"\nHausman Test: χ²={hausman_res['stat']:.4f}, p={hausman_res['p']:.4f} (df={hausman_res['df']})"
        bp_str = ""
        if stats.get("BP_stat") is not None:
            bp_str = f"\nBreusch-Pagan Test: LM={stats['BP_stat']:.4f}, p={stats['BP_p']:.4f}"
        re_str = ""
        if st.session_state.model_type == "Random Effects (GLS)":
            re_str = f"\nVariance Components: σ²ᵤ={stats.get('sigma_u2',0):.6f}, σ²ₑ={stats.get('sigma_e2',0):.6f}, θ={stats.get('theta',0):.4f}"
        context = f"""
Model: {st.session_state.model_type}
Dependent variable: {res['y_col']}
Independent variables: {', '.join(res['x_cols'])}
Entity column: {res['entity_col']} | Time column: {res['time_col']}

Fit Statistics:
  R²        = {stats['R2']:.4f}
  Adj. R²   = {stats['R2_adj']:.4f}
  N         = {stats['N']}
  AIC       = {stats['AIC']:.2f}
  BIC       = {stats['BIC']:.2f}
  F-stat    = {stats.get('F_stat', 'N/A')}  (p = {stats.get('F_p', 'N/A')}){hausman_str}{bp_str}{re_str}

Coefficient Table:
{coeff_table}
"""

        sys_prompt = (
            "You are an expert econometrician. Given panel regression output, produce a structured, "
            "publication-quality interpretation that will be rendered into a Word document. "
            "Format your response EXACTLY as follows, using these markers:\n\n"
            "## Model Specification\n"
            "Brief paragraph on why this estimator is appropriate, citing the Hausman test result if available.\n\n"
            "## Coefficient Interpretation\n"
            "For each variable, use a bullet: '- **VariableName**: [coeff value], [SE], [t-stat], p=[p-value]. "
            "Interpretation: [one sentence on economic/practical meaning and direction of effect].'\n\n"
            "## Model Fit & Overall Significance\n"
            "Discuss R², Adjusted R², F-statistic and its p-value, AIC/BIC in one paragraph. "
            "Use **bold** for specific numeric values.\n\n"
            "## Diagnostic Test Results\n"
            "- **Jarque-Bera**: state stat and p, conclude on residual normality.\n"
            "- **Breusch-Pagan**: state LM stat and p, conclude on heteroskedasticity.\n"
            "- **Durbin-Watson**: state value, conclude on autocorrelation.\n\n"
            "## Caveats & Concerns\n"
            "Numbered list of up to 3 potential issues (endogeneity, measurement error, etc.).\n\n"
            "## Recommendations\n"
            "Numbered list of up to 3 concrete next steps for the analyst.\n\n"
            "Use **bold** for key terms and statistics. Keep each section concise. "
            "Do NOT use plain prose without structure. Do NOT add preamble or closing remarks."
        )

        col_explain, col_custom = st.columns([3, 2])

        with col_explain:
            if st.button("✦ Generate AI Explanation", type="primary", width='stretch'):
                with st.spinner("AI model is analysing your results…"):
                    explanation = call_openai(sys_prompt, f"Please explain these panel regression results:\n\n{context}")
                    st.session_state.ai_explanation = explanation

        with col_custom:
            custom_q = st.text_input("Ask a specific question about the results…",
                                      placeholder="e.g. Is x1 economically significant?")
            if st.button("Ask AI Model", width='stretch') and custom_q:
                with st.spinner("Thinking…"):
                    answer = call_openai(
                        sys_prompt,
                        f"Here are the regression results:\n\n{context}\n\nQuestion: {custom_q}",
                    )
                    st.session_state.ai_explanation = answer

        if st.session_state.ai_explanation:
            st.markdown("---")
            st.markdown(f"""
            <div class="ai-label">✦ &nbsp; AI MODEL INTERPRETATION</div>
            <div class="ai-box">{st.session_state.ai_explanation}</div>
            """, unsafe_allow_html=True)

        # ── Quick insight cards ────────────────────────────────────────────────
        st.markdown("---")
        st.markdown('<div class="scard-title">Quick Insights</div>', unsafe_allow_html=True)

        ic1, ic2, ic3 = st.columns(3)
        sig_vars = result_df[result_df["p_value"] < 0.05]["Variable"].tolist()
        insig_vars = result_df[result_df["p_value"] >= 0.05]["Variable"].tolist()
        r2_val = stats["R2"]

        with ic1:
            st.markdown(f"""
            <div class="scard">
                <div class="scard-title">Significance</div>
                <div style="color:var(--accent);font-size:1.2rem;font-family:'Syne',sans-serif;font-weight:700;">{len(sig_vars)}/{len(result_df)}</div>
                <div style="color:var(--muted);font-size:0.75rem;margin-top:4px;">variables significant at 5%</div>
                <div style="margin-top:10px;font-size:0.72rem;color:var(--text);">{', '.join(sig_vars) if sig_vars else '—'}</div>
            </div>
            """, unsafe_allow_html=True)
        with ic2:
            r2_color = "#00e5c8" if r2_val > 0.7 else "#f5a623" if r2_val > 0.4 else "#f05c7c"
            r2_label = "Strong fit" if r2_val > 0.7 else "Moderate fit" if r2_val > 0.4 else "Weak fit"
            st.markdown(f"""
            <div class="scard">
                <div class="scard-title">Model Fit</div>
                <div style="color:{r2_color};font-size:1.2rem;font-family:'Syne',sans-serif;font-weight:700;">{r2_val:.4f}</div>
                <div style="color:var(--muted);font-size:0.75rem;margin-top:4px;">R-squared · {r2_label}</div>
            </div>
            """, unsafe_allow_html=True)
        with ic3:
            largest = result_df.iloc[result_df["Coeff"].abs().argmax()]
            st.markdown(f"""
            <div class="scard">
                <div class="scard-title">Largest Effect</div>
                <div style="color:var(--accent2);font-size:1.2rem;font-family:'Syne',sans-serif;font-weight:700;">{largest['Variable']}</div>
                <div style="color:var(--muted);font-size:0.75rem;margin-top:4px;">coeff = {largest['Coeff']:.4f}</div>
            </div>
            """, unsafe_allow_html=True)


# ─── Footer ───────────────────────────────────────────────────────────────────
st.markdown("---")
st.markdown("""
<div style="text-align:center;font-family:'DM Mono',monospace;font-size:0.7rem;color:var(--muted);padding:12px 0;">
    ⬡ PanelStatX · Panel Regression Analysis System · Powered by an AI Model
</div>
""", unsafe_allow_html=True)
